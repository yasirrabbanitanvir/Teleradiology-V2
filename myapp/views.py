from rest_framework import viewsets, status, generics
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.authtoken.models import Token
from rest_framework.authentication import TokenAuthentication
from rest_framework.permissions import IsAuthenticated, BasePermission, AllowAny
from rest_framework.decorators import (
    action,
    api_view,
    permission_classes,
    authentication_classes,
)
from rest_framework.pagination import PageNumberPagination
from django.contrib.auth import authenticate
from django.contrib.auth.models import User, Group
from django.contrib.auth.hashers import make_password
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from django.views.decorators.http import require_http_methods
from django.http import FileResponse, Http404, JsonResponse, HttpResponse
from django.shortcuts import render
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.db.models import Q, Count, Sum, Max
from django.core.cache import cache
from django_filters import rest_framework as filters
from pathlib import Path
import os
import mimetypes
from datetime import datetime, timezone
import logging
import json
import pydicom
from pydicom.dataset import FileMetaDataset, Dataset
from pydicom.uid import (
    ExplicitVRLittleEndian,
    SecondaryCaptureImageStorage,
    generate_uid,
)
from pydicom.encaps import encapsulate
import traceback
import tempfile
from django.core.signals import request_finished
from .models import (
    Patient,
    DoctorUpload,
    UserProfile,
    Center,
    CenterName,
    DICOMImage,
    ReportTemplate,
    UserSession,
)
from .serializers import (
    PatientSerializer,
    DoctorUploadSerializer,
    UserSerializer,
    CenterSerializer,
    ReportSerializer,
    DICOMImageSerializer,
    ReportTemplateSerializer,
)
from PIL import Image
import numpy as np
import io
from rest_framework.authtoken.models import Token

logger = logging.getLogger(__name__)


def _get_history_path(dicom_id):
    from django.conf import settings
    history_dir = os.path.join(getattr(settings, 'MEDIA_ROOT', ''), 'report_history')
    os.makedirs(history_dir, exist_ok=True)
    return os.path.join(history_dir, f'history_{dicom_id}.json')

def _read_history(dicom_id):
    try:
        path = _get_history_path(dicom_id)
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception:
        pass
    return []

def _write_history(dicom_id, entries):
    try:
        with open(_get_history_path(dicom_id), 'w', encoding='utf-8') as f:
            json.dump(entries, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning(f"Could not write report history for {dicom_id}: {e}")

def _append_to_history(dicom_id, entry):
    entries = _read_history(dicom_id)
    entries.insert(0, entry)  
    _write_history(dicom_id, entries)


def _build_center_name_to_institute_map():
    cache_key = "center_name_institute_map_v1"
    cached = cache.get(cache_key)
    if cached is not None:
        return cached
    mapping = {}
    for cn in CenterName.objects.select_related('center').all():
        if cn.center_id:
            mapping[cn.name] = cn.center.institute_name or cn.name
    for center in Center.objects.all():
        mapping.setdefault(center.institute_name, center.institute_name)
    cache.set(cache_key, mapping, 30)
    return mapping



def decompress_dicom(ds):
    try:
        if hasattr(ds, "file_meta") and hasattr(ds.file_meta, "TransferSyntaxUID"):
            transfer_syntax = str(ds.file_meta.TransferSyntaxUID)
            compressed_syntaxes = [
                "1.2.840.10008.1.2.4.50",
                "1.2.840.10008.1.2.4.51",
                "1.2.840.10008.1.2.4.57",
                "1.2.840.10008.1.2.4.70",
                "1.2.840.10008.1.2.4.80",
                "1.2.840.10008.1.2.4.81",
                "1.2.840.10008.1.2.4.90",
                "1.2.840.10008.1.2.4.91",
                "1.2.840.10008.1.2.5",
            ]
            if transfer_syntax in compressed_syntaxes:
                logger.info(
                    f"Decompressing DICOM with transfer syntax: {transfer_syntax}"
                )
                ds.decompress()
                ds.file_meta.TransferSyntaxUID = pydicom.uid.ExplicitVRLittleEndian
                return True
        return False
    except Exception as e:
        logger.error(f"Error decompressing DICOM: {str(e)}")
        raise


def serve_dicom(request, filename):
    file_path = os.path.join(settings.MEDIA_ROOT, filename)
    if not os.path.exists(file_path):
        raise Http404("DICOM file not found")
    try:
        ds = pydicom.dcmread(file_path, force=True)
        was_decompressed = decompress_dicom(ds)
        if was_decompressed:
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".dcm")
            ds.save_as(temp_file.name, write_like_original=False)
            temp_file.close()
            response = FileResponse(
                open(temp_file.name, "rb"), content_type="application/dicom"
            )

            def cleanup_temp_file(sender, **kwargs):
                try:
                    os.unlink(temp_file.name)
                except:
                    pass

            request_finished.connect(cleanup_temp_file, weak=False)
        else:
            response = FileResponse(
                open(file_path, "rb"), content_type="application/dicom"
            )
        response["Access-Control-Allow-Origin"] = "*"
        response["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        response["Access-Control-Max-Age"] = "1000"
        response["Access-Control-Allow-Headers"] = "X-Requested-With, Content-Type"
        return response
    except Exception as e:
        logger.error(f"Error serving DICOM file: {str(e)}")
        response = FileResponse(open(file_path, "rb"), content_type="application/dicom")
        response["Access-Control-Allow-Origin"] = "*"
        response["Access-Control-Allow-Methods"] = "GET, OPTIONS"
        response["Access-Control-Max-Age"] = "1000"
        response["Access-Control-Allow-Headers"] = "X-Requested-With, Content-Type"
        return response


class RoleBasedPermission(BasePermission):
    allowed_roles = []

    def has_permission(self, request, view):
        if not request.user.is_authenticated:
            return False
        profile = UserProfile.objects.filter(user=request.user).first()
        if not profile:
            return request.user.is_superuser
        if not profile.role:
            return request.user.is_superuser
        return profile.role.name in self.allowed_roles


class AdminPermission(RoleBasedPermission):
    allowed_roles = ["Admin"]


class DoctorPermission(RoleBasedPermission):
    allowed_roles = ["Doctor"]


class SubAdminPermission(RoleBasedPermission):
    allowed_roles = ["SubAdmin"]


class CenterPermission(RoleBasedPermission):
    allowed_roles = ["Center"]


class AdminOrSubAdminPermission(RoleBasedPermission):
    allowed_roles = ["Admin", "SubAdmin"]


class PatientViewSet(viewsets.ModelViewSet):
    queryset = Patient.objects.all()
    serializer_class = PatientSerializer
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated]

    @action(detail=False, methods=["get"])
    def by_patient_id(self, request):
        patient_id = request.query_params.get("patient_id")
        if patient_id:
            patient = Patient.objects.filter(patient_id=patient_id).first()
            if patient:
                serializer = self.get_serializer(patient)
                return Response(serializer.data)
        return Response(
            {"detail": "Patient not found"}, status=status.HTTP_404_NOT_FOUND
        )


@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def get_user_contact_number(request):
    try:
        user_profile = UserProfile.objects.filter(user=request.user).first()
        if user_profile and user_profile.contact_number:
            return Response({'contact_number': user_profile.contact_number})
        return Response({'contact_number': 'N/A'})
    except Exception as e:
        logger.error(f"Error fetching user contact number: {str(e)}")
        return Response({'contact_number': 'N/A'})


class DoctorUploadViewSet(viewsets.ModelViewSet):
    queryset = DoctorUpload.objects.all()
    serializer_class = DoctorUploadSerializer
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated, DoctorPermission | CenterPermission]

    def get_queryset(self):
        queryset = super().get_queryset()
        patient_id = self.request.query_params.get("patient")
        if patient_id:
            queryset = queryset.filter(patient_id=patient_id)
        return queryset


@method_decorator(csrf_exempt, name="dispatch")
class CustomLoginView(APIView):
    def post(self, request):
        username = request.data.get("username")
        password = request.data.get("password")
        user = authenticate(request, username=username, password=password)
        if user:
            token, created = Token.objects.get_or_create(user=user)
            from django.utils import timezone
            x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
            if x_forwarded_for:
                client_ip = x_forwarded_for.split(',')[0].strip()
            else:
                client_ip = request.META.get('REMOTE_ADDR')
            now = timezone.now()
            user.last_login = now
            user.save(update_fields=['last_login'])
            UserSession.objects.update_or_create(
                user=user,
                defaults={'login_time': now, 'last_activity': now, 'client_ip': client_ip}
            )
            profile = UserProfile.objects.filter(user=user).first()
            role = (
                profile.role.name
                if profile and profile.role
                else ("Admin" if user.is_superuser else "Unknown")
            )
            redirect_url = "/"
            center_name = None
            institute_name = None
            if role == "Admin":
                redirect_url = "/admin/"
            elif role == "SubAdmin":
                redirect_url = "/static/index.html"
            elif role == "Center":
                redirect_url = "/static/institute.html"
                if profile and profile.center:
                    center_obj = profile.center
                    institute_name = center_obj.institute_name
                    center_name_obj = center_obj.center_names.first()
                    center_name = center_name_obj.name if center_name_obj else None
                else:
                    center = Center.objects.filter(user=user).first()
                    if center:
                        institute_name = center.institute_name
                        center_name_obj = center.center_names.first()
                        center_name = center_name_obj.name if center_name_obj else None
                        if profile:
                            profile.center = center
                            profile.save()
            elif role == "Doctor":
                redirect_url = "/static/doctor.html"
            return Response(
                {
                    "token": token.key,
                    "redirect": redirect_url,
                    "role": role,
                    "center_name": center_name,
                    "institute_name": institute_name,
                }
            )
        return Response({"error": "Invalid credentials"}, status=401)


class UserViewSet(viewsets.ViewSet):
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated, AdminPermission]

    def list(self, request):
        users = User.objects.all().select_related('userprofile__role', 'userprofile__center')
        for user in users:
            if not hasattr(user, 'userprofile') or user.userprofile is None:
                try:
                    UserProfile.objects.get_or_create(user=user)
                except Exception:
                    pass
        users = User.objects.all().select_related('userprofile__role', 'userprofile__center')
        serializer = UserSerializer(users, many=True, context={"request": request})
        return Response(serializer.data)

    def create(self, request):
        username = request.data.get("username")
        password = request.data.get("password")
        role = request.data.get("role")
        center_id = request.data.get("center")
        if User.objects.filter(username=username).exists():
            return Response(
                {"detail": "Username already exists"},
                status=status.HTTP_400_BAD_REQUEST,
            )
        user = User.objects.create(username=username, password=make_password(password))
        group = Group.objects.get(name=role)
        profile = UserProfile.objects.create(user=user, role=group)
        if role == "Center" and center_id:
            center = Center.objects.get(id=center_id)
            profile.center = center
            profile.save()
            center.user = user
            center.save()
        return Response(
            {"detail": "User created successfully"}, status=status.HTTP_201_CREATED
        )

    def destroy(self, request, pk=None):
        try:
            user = User.objects.get(pk=pk)
            user.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except User.DoesNotExist:
            return Response(
                {"detail": "User not found"}, status=status.HTTP_404_NOT_FOUND
            )


class CenterViewSet(viewsets.ModelViewSet):
    queryset = Center.objects.all()
    serializer_class = CenterSerializer
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated, AdminPermission]
    
    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        centers_data = []
        
        for center in queryset:
            assigned_doctors = center.assigned_doctors.all()
            doctor_names = [
                doc.full_name or doc.user.username 
                for doc in assigned_doctors
            ]
            
            serializer = self.get_serializer(center)
            center_data = serializer.data
            center_data['assigned_doctors_names'] = doctor_names
            centers_data.append(center_data)
        
        return Response(centers_data)


def extract_dicom_metadata(ds):
    metadata = {}

    def safe_get_attr(ds, attr, default=""):
        try:
            if not hasattr(ds, attr):
                return default
            value = getattr(ds, attr)
            if value is None:
                return default
            if hasattr(value, "original_string"):
                try:
                    return str(value)
                except UnicodeDecodeError:
                    for encoding in ["latin-1", "iso-8859-1", "cp1252", "ascii"]:
                        try:
                            if isinstance(value.original_string, bytes):
                                return value.original_string.decode(
                                    encoding, errors="replace"
                                )
                            return str(value)
                        except:
                            continue
                    return default
                except Exception:
                    return default
            if isinstance(value, (list, tuple)):
                try:
                    return [str(item) for item in value]
                except:
                    return default
            if hasattr(value, "alphabetic"):
                try:
                    return str(value.alphabetic) if value.alphabetic else default
                except UnicodeDecodeError:
                    return (
                        value.alphabetic.decode("latin-1", errors="replace")
                        if value.alphabetic
                        else default
                    )
                except:
                    return default
            if isinstance(value, bytes):
                if attr in [
                    "PatientName",
                    "PatientID",
                    "StudyDescription",
                    "SeriesDescription",
                ]:
                    for encoding in [
                        "utf-8",
                        "latin-1",
                        "iso-8859-1",
                        "cp1252",
                        "ascii",
                    ]:
                        try:
                            return value.decode(encoding, errors="replace")
                        except UnicodeDecodeError:
                            continue
                    return default
                else:
                    return default
            try:
                return str(value)
            except UnicodeDecodeError:
                if hasattr(value, "encode"):
                    return value.encode("latin-1", errors="replace").decode("latin-1")
                return default
        except Exception as e:
            logger.warning(f"Error extracting {attr}: {str(e)}")
            return default

    metadata["patient_name"] = safe_get_attr(ds, "PatientName", "")
    metadata["patient_id"] = safe_get_attr(ds, "PatientID", "")
    metadata["patient_birth_date"] = safe_get_attr(ds, "PatientBirthDate", "")
    metadata["patient_sex"] = safe_get_attr(ds, "PatientSex", "")
    metadata["study_instance_uid"] = safe_get_attr(ds, "StudyInstanceUID", "")
    metadata["study_date"] = safe_get_attr(ds, "StudyDate", "")
    metadata["study_time"] = safe_get_attr(ds, "StudyTime", "")
    metadata["study_description"] = safe_get_attr(ds, "StudyDescription", "")
    metadata["referring_physician"] = safe_get_attr(ds, "ReferringPhysicianName", "")
    metadata["series_instance_uid"] = safe_get_attr(ds, "SeriesInstanceUID", "")
    metadata["series_number"] = safe_get_attr(ds, "SeriesNumber", "")
    metadata["series_description"] = safe_get_attr(ds, "SeriesDescription", "")
    metadata["modality"] = safe_get_attr(ds, "Modality", "")
    metadata["sop_instance_uid"] = safe_get_attr(ds, "SOPInstanceUID", "")
    metadata["instance_number"] = safe_get_attr(ds, "InstanceNumber", "")
    try:
        if hasattr(ds, "ImageOrientationPatient") and ds.ImageOrientationPatient:
            orient = ds.ImageOrientationPatient
            if hasattr(orient, "__iter__"):
                metadata["image_orientation"] = [float(x) for x in orient]
            else:
                metadata["image_orientation"] = str(orient)
    except:
        metadata["image_orientation"] = ""
    try:
        if hasattr(ds, "ImagePositionPatient") and ds.ImagePositionPatient:
            pos = ds.ImagePositionPatient
            if hasattr(pos, "__iter__"):
                metadata["image_position"] = [float(x) for x in pos]
            else:
                metadata["image_position"] = str(pos)
    except:
        metadata["image_position"] = ""
    try:
        if hasattr(ds, "PixelSpacing") and ds.PixelSpacing:
            spacing = ds.PixelSpacing
            if hasattr(spacing, "__iter__"):
                metadata["pixel_spacing"] = [float(x) for x in spacing]
            else:
                metadata["pixel_spacing"] = str(spacing)
    except:
        metadata["pixel_spacing"] = ""
    try:
        if hasattr(ds, "SliceThickness") and ds.SliceThickness:
            metadata["slice_thickness"] = float(ds.SliceThickness)
    except:
        metadata["slice_thickness"] = None
    return metadata


def convert_image_bytes_to_dicom(
    image_bytes, original_filename=None, provided_meta=None, prefer_transfer_syntax=None
):
    img = Image.open(io.BytesIO(image_bytes))
    filename_lower = (original_filename or "").lower()
    image_is_jpeg_file = filename_lower.endswith(".jpg") or filename_lower.endswith(
        ".jpeg"
    )
    want_jpeg_encapsulation = image_is_jpeg_file and (
        prefer_transfer_syntax in (None, "auto", "jpegbaseline")
        or (prefer_transfer_syntax and "jpeg" in prefer_transfer_syntax.lower())
    )
    file_meta = FileMetaDataset()
    file_meta.MediaStorageSOPClassUID = SecondaryCaptureImageStorage
    file_meta.MediaStorageSOPInstanceUID = generate_uid()
    file_meta.TransferSyntaxUID = ExplicitVRLittleEndian
    file_meta.ImplementationClassUID = generate_uid()
    ds = Dataset()
    ds.file_meta = file_meta
    ds.SOPClassUID = file_meta.MediaStorageSOPClassUID
    ds.SOPInstanceUID = file_meta.MediaStorageSOPInstanceUID
    if provided_meta and provided_meta.get("study_instance_uid"):
        ds.StudyInstanceUID = provided_meta.get("study_instance_uid")
    else:
        ds.StudyInstanceUID = generate_uid()
    if provided_meta and provided_meta.get("series_instance_uid"):
        ds.SeriesInstanceUID = provided_meta.get("series_instance_uid")
    else:
        ds.SeriesInstanceUID = generate_uid()
    ds.PatientName = provided_meta.get("patient_name", "") if provided_meta else ""
    ds.PatientID = provided_meta.get("patient_id", "") if provided_meta else ""
    ds.PatientSex = provided_meta.get("patient_sex", "") if provided_meta else ""
    ds.PatientBirthDate = provided_meta.get("patient_birth_date", "") if provided_meta else ""
    ds.StudyDate = (
        provided_meta.get("study_date", datetime.now().strftime("%Y%m%d"))
        if provided_meta
        else datetime.now().strftime("%Y%m%d")
    )
    ds.StudyTime = (
        provided_meta.get("study_time", datetime.now().strftime("%H%M%S"))
        if provided_meta
        else datetime.now().strftime("%H%M%S")
    )
    ds.Modality = provided_meta.get("modality", "OT") if provided_meta else "OT"
    ds.StudyDescription = (
        provided_meta.get("study_description", "") if provided_meta else ""
    )
    ds.SeriesDescription = (
        provided_meta.get("series_description", "") if provided_meta else ""
    )
    ds.ReferringPhysicianName = (
        provided_meta.get("referring_physician", "") if provided_meta else ""
    )
    ds.InstitutionName = (
        provided_meta.get("institution_name", "") if provided_meta else ""
    )
    if provided_meta and provided_meta.get("instance_number"):
        ds.InstanceNumber = str(provided_meta.get("instance_number"))
    else:
        ds.InstanceNumber = "1"
    if provided_meta and provided_meta.get("series_number"):
        ds.SeriesNumber = str(provided_meta.get("series_number"))
    else:
        ds.SeriesNumber = "1"
    if want_jpeg_encapsulation:
        try:
            try:
                img.load()
            except Exception:
                pass
            jpeg_bytes = image_bytes
            if img.mode in ("RGB", "CMYK"):
                ds.SamplesPerPixel = 3
                ds.PhotometricInterpretation = "YBR_FULL"
                ds.PlanarConfiguration = 0
            else:
                ds.SamplesPerPixel = 1
                ds.PhotometricInterpretation = "MONOCHROME2"
                ds.PlanarConfiguration = 0
            ds.Rows = img.size[1]
            ds.Columns = img.size[0]
            ds.BitsAllocated = 8
            ds.BitsStored = 8
            ds.HighBit = 7
            ds.PixelRepresentation = 0
            if prefer_transfer_syntax and "lossless" in prefer_transfer_syntax.lower():
                try:
                    file_meta.TransferSyntaxUID = pydicom.uid.JPEGLossless
                    try:
                        import imagecodecs

                        arr = np.asarray(img)
                        if hasattr(imagecodecs, "jpeg_lossless_encode"):
                            jpeg_fragment = imagecodecs.jpeg_lossless_encode(arr)
                        elif hasattr(imagecodecs, "jpeg_ls_encode"):
                            jpeg_fragment = imagecodecs.jpeg_ls_encode(arr)
                        else:
                            raise Exception("imagecodecs lossless encode api not found")
                        ds.PixelData = encapsulate([jpeg_fragment])
                    except Exception as e:
                        logger.warning(f"JPEG Lossless encoding failed: {e}")
                        ds.PixelData = encapsulate([jpeg_bytes])
                except Exception:
                    file_meta.TransferSyntaxUID = pydicom.uid.JPEGBaseline8Bit
                    ds.PixelData = encapsulate([jpeg_bytes])
            else:
                file_meta.TransferSyntaxUID = pydicom.uid.JPEGBaseline8Bit
                ds.PixelData = encapsulate([jpeg_bytes])
            ds.is_little_endian = True
            ds.is_implicit_VR = False
            buf = io.BytesIO()
            pydicom.filewriter.dcmwrite(buf, ds, write_like_original=False)
            dcm_bytes = buf.getvalue()
            buf.close()
            ds_read = pydicom.dcmread(
                pydicom.filebase.DicomBytesIO(dcm_bytes), force=True
            )
            return dcm_bytes, ds_read
        except Exception as e:
            logger.warning(f"Encapsulation failed, falling back: {e}")
    try:
        if img.mode not in ("L", "I;16", "I", "RGB"):
            try:
                if "A" in img.mode or img.mode == "RGBA":
                    img = img.convert("RGB")
                else:
                    img = img.convert("L")
            except Exception:
                img = img.convert("L")
        img.load()
        pixel_array = np.asarray(img)
        if pixel_array.ndim == 3:
            samples_per_pixel = pixel_array.shape[2]
            if samples_per_pixel == 3:
                ds.SamplesPerPixel = 3
                ds.PhotometricInterpretation = "RGB"
                ds.PlanarConfiguration = 0
            elif samples_per_pixel == 4:
                pixel_array = pixel_array[:, :, :3]
                ds.SamplesPerPixel = 3
                ds.PhotometricInterpretation = "RGB"
                ds.PlanarConfiguration = 0
            else:
                pixel_array = pixel_array.mean(axis=2).astype(np.uint8)
                ds.SamplesPerPixel = 1
                ds.PhotometricInterpretation = "MONOCHROME2"
        else:
            ds.SamplesPerPixel = 1
            ds.PhotometricInterpretation = "MONOCHROME2"
            if pixel_array.ndim != 2:
                pixel_array = pixel_array.reshape(pixel_array.shape[0], -1)
        if pixel_array.dtype != np.uint8 and pixel_array.dtype != np.uint16:
            pmin = pixel_array.min()
            pmax = pixel_array.max()
            if pmax > pmin:
                scaled = (pixel_array.astype(np.float32) - float(pmin)) / float(
                    pmax - pmin
                )
                pixel_array = (scaled * 255.0).astype(np.uint8)
            else:
                pixel_array = (pixel_array * 0).astype(np.uint8)
        rows, cols = pixel_array.shape[:2]
        ds.Rows = int(rows)
        ds.Columns = int(cols)
        if pixel_array.dtype == np.uint16:
            ds.BitsAllocated = 16
            ds.BitsStored = 16
            ds.HighBit = 15
            ds.PixelRepresentation = 0
        else:
            ds.BitsAllocated = 8
            ds.BitsStored = 8
            ds.HighBit = 7
            ds.PixelRepresentation = 0
        ds.PixelData = pixel_array.tobytes()
        file_meta.TransferSyntaxUID = ExplicitVRLittleEndian
        ds.is_little_endian = True
        ds.is_implicit_VR = False
        if original_filename:
            try:
                ds.ImageComments = f"SourceFile: {original_filename}"
            except:
                pass
        buf = io.BytesIO()
        pydicom.filewriter.dcmwrite(buf, ds, write_like_original=False)
        dcm_bytes = buf.getvalue()
        buf.close()
        ds_read = pydicom.dcmread(pydicom.filebase.DicomBytesIO(dcm_bytes), force=True)
        return dcm_bytes, ds_read
    except Exception as e:
        logger.error(f"Failed to convert image: {e}")
        raise


def ensure_decompressed_and_normalized(ds):
    try:
        arr = ds.pixel_array
    except Exception as e:
        logger.warning(f"Could not decode pixel data: {e}")
        return None, ds
    try:
        new_ds = ds.copy()
        file_meta = FileMetaDataset()
        file_meta.MediaStorageSOPClassUID = getattr(
            new_ds, "SOPClassUID", SecondaryCaptureImageStorage
        )
        file_meta.MediaStorageSOPInstanceUID = getattr(
            new_ds, "SOPInstanceUID", generate_uid()
        )
        file_meta.TransferSyntaxUID = ExplicitVRLittleEndian
        file_meta.ImplementationClassUID = generate_uid()
        new_ds.file_meta = file_meta
        pixel_array = np.asarray(arr)
        if pixel_array.ndim == 4:
            n_frames, rows, cols, samples = pixel_array.shape
            new_ds.NumberOfFrames = str(n_frames)
            frames_flat = pixel_array
        elif pixel_array.ndim == 3:
            if pixel_array.shape[2] in (1, 3, 4):
                rows, cols, samples = pixel_array.shape
                new_ds.NumberOfFrames = "1"
                frames_flat = pixel_array.reshape((1, rows, cols, samples))
            else:
                n_frames, rows, cols = pixel_array.shape
                samples = 1
                new_ds.NumberOfFrames = str(n_frames)
                frames_flat = pixel_array.reshape((n_frames, rows, cols, 1))
        elif pixel_array.ndim == 2:
            rows, cols = pixel_array.shape
            samples = 1
            new_ds.NumberOfFrames = "1"
            frames_flat = pixel_array.reshape((1, rows, cols, 1))
        else:
            frames_flat = pixel_array.reshape((1,) + pixel_array.shape)
            rows = int(frames_flat.shape[1])
            cols = int(frames_flat.shape[2])
            samples = int(frames_flat.shape[3]) if frames_flat.ndim >= 4 else 1
            new_ds.NumberOfFrames = str(frames_flat.shape[0])
        new_ds.SamplesPerPixel = samples
        if samples == 1:
            new_ds.PhotometricInterpretation = getattr(
                new_ds, "PhotometricInterpretation", "MONOCHROME2"
            )
        elif samples == 3:
            new_ds.PhotometricInterpretation = "RGB"
            new_ds.PlanarConfiguration = 0
        else:
            new_ds.PhotometricInterpretation = getattr(
                new_ds, "PhotometricInterpretation", "RGB"
            )
            new_ds.PlanarConfiguration = getattr(new_ds, "PlanarConfiguration", 0)
        new_ds.Rows = int(rows)
        new_ds.Columns = int(cols)
        if frames_flat.dtype == np.uint16:
            new_ds.BitsAllocated = 16
            new_ds.BitsStored = 16
            new_ds.HighBit = 15
            new_ds.PixelRepresentation = 0
        else:
            new_ds.BitsAllocated = 8
            new_ds.BitsStored = 8
            new_ds.HighBit = 7
            new_ds.PixelRepresentation = 0
        new_ds.PixelData = frames_flat.tobytes()
        new_ds.is_little_endian = True
        new_ds.is_implicit_VR = False
        buf = io.BytesIO()
        pydicom.filewriter.dcmwrite(buf, new_ds, write_like_original=False)
        dcm_bytes = buf.getvalue()
        buf.close()
        ds_read = pydicom.dcmread(pydicom.filebase.DicomBytesIO(dcm_bytes), force=True)
        return dcm_bytes, ds_read
    except Exception as ex:
        logger.warning(f"Failed to normalize: {ex}")
        return None, ds

@csrf_exempt
def receive_dicom_data(request):
    try:
        logger.info("Received DICOM data request")
        if request.method == "OPTIONS":
            response = HttpResponse()
            response["Access-Control-Allow-Origin"] = "*"
            response["Access-Control-Allow-Methods"] = "POST, OPTIONS"
            response["Access-Control-Allow-Headers"] = "Content-Type"
            return response
        if request.method != "POST":
            return JsonResponse(
                {"success": False, "error": "Only POST method allowed"}, status=405
            )
        if "dicom_file" not in request.FILES:
            logger.error("No DICOM file in request")
            return JsonResponse(
                {"success": False, "error": "No DICOM file provided"}, status=400
            )
        center_name = request.POST.get("center_name", "").strip()
        if not center_name:
            logger.error("No center name provided")
            return JsonResponse(
                {"success": False, "error": "Center name is required"}, status=400
            )
        uploaded = request.FILES["dicom_file"]
        filename = uploaded.name or "upload"
        filename_lower = filename.lower()
        file_content = uploaded.read()
        logger.info(f"Processing file: {filename}, size: {len(file_content)} bytes")
        image_exts = (".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff")
        ds = None
        is_converted_from_image = False
        if any(filename_lower.endswith(ext) for ext in image_exts):
            logger.info(f"File {filename} is an image, converting to DICOM")
            try:
                patient_age = request.POST.get("patient_age", "")
                patient_birth_date = ""
                if patient_age and patient_age.strip() not in ("", "N/A", "n/a") and patient_age.replace(" ", "").isdigit():
                    current_year = datetime.now().year
                    birth_year = current_year - int(patient_age.strip())
                    patient_birth_date = f"{birth_year}0101"
                
                provided_meta = {
                    "patient_name": request.POST.get("patient_name", ""),
                    "patient_id": request.POST.get("patient_id", ""),
                    "patient_sex": request.POST.get("patient_sex", ""),
                    "patient_birth_date": patient_birth_date,
                    "study_date": request.POST.get("study_date", ""),
                    "study_time": request.POST.get("study_time", ""),
                    "modality": request.POST.get("modality", "OT"),
                    "study_description": request.POST.get("study_description", ""),
                    "series_description": request.POST.get("series_description", ""),
                    "referring_physician": request.POST.get("referring_physician", ""),
                    "study_instance_uid": request.POST.get("study_instance_uid", ""),
                    "series_instance_uid": request.POST.get("series_instance_uid", ""),
                    "instance_number": request.POST.get("instance_number", "1"),
                    "series_number": request.POST.get("series_number", "1"),
                }
                converted_dcm_bytes, ds = convert_image_bytes_to_dicom(
                    file_content,
                    original_filename=filename,
                    provided_meta=provided_meta,
                )
                file_content = converted_dcm_bytes
                is_converted_from_image = True
                logger.info(f"Successfully converted image {filename} to DICOM")
            except Exception as e:
                logger.error(f"Image conversion failed: {str(e)}")
                traceback.print_exc()
                return JsonResponse(
                    {"success": False, "error": f"Image conversion failed: {str(e)}"},
                    status=400,
                )
        else:
            logger.info(f"Processing as DICOM file: {filename}")
            try:
                ds = pydicom.dcmread(
                    pydicom.filebase.DicomBytesIO(file_content), force=True
                )
                if not hasattr(ds, "SpecificCharacterSet"):
                    ds.SpecificCharacterSet = "ISO_IR 100"
                logger.info(f"Successfully read DICOM file: {filename}")
            except Exception as e:
                logger.error(f"Error reading DICOM file {filename}: {str(e)}")
                traceback.print_exc()
                return JsonResponse(
                    {
                        "success": False, 
                        "error": f"Invalid DICOM file: {str(e)}. Please ensure you're uploading a valid DICOM (.dcm) file or supported image format."
                    }, 
                    status=400
                )
        try:
            metadata = extract_dicom_metadata(ds)
            logger.info(f"Extracted metadata: Patient ID={metadata.get('patient_id', 'N/A')}")
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            traceback.print_exc()
            return JsonResponse(
                {"success": False, "error": f"Metadata extraction error: {str(e)}"}, 
                status=500
            )
        try:
            center_dir = (
                center_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
            )
            sop_uid = metadata.get('sop_instance_uid', 'unknown')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            out_filename = f"{center_dir}/{sop_uid}_{timestamp}.dcm"
            file_path = default_storage.save(out_filename, ContentFile(file_content))
            full_path = os.path.join(settings.MEDIA_ROOT, file_path)
            logger.info(f"File saved to: {full_path}")
        except Exception as e:
            logger.error(f"Error saving file: {str(e)}")
            traceback.print_exc()
            return JsonResponse(
                {"success": False, "error": f"File save error: {str(e)}"}, 
                status=500
            )
        try:
            dicom_image = DICOMImage.objects.create(
                center_name=center_name[:200],
                patient_name=metadata.get("patient_name", "")[:200],
                patient_id=metadata.get("patient_id", "")[:64],
                patient_birth_date=metadata.get("patient_birth_date", "") or None,
                patient_sex=metadata.get("patient_sex", "")[:10],
                study_instance_uid=metadata.get("study_instance_uid", "")[:64],
                study_date=metadata.get("study_date", "") or None,
                study_time=metadata.get("study_time", "") or None,
                study_description=metadata.get("study_description", "")[:200],
                referring_physician=metadata.get("referring_physician", "")[:200],
                series_instance_uid=metadata.get("series_instance_uid", "")[:64],
                series_number=metadata.get("series_number", "") or None,
                series_description=metadata.get("series_description", "")[:200],
                modality=metadata.get("modality", "")[:16],
                sop_instance_uid=metadata.get("sop_instance_uid", "")[:64],
                instance_number=metadata.get("instance_number", "") or None,
                file_path=file_path,
                file_size=len(file_content),
                image_orientation=(
                    json.dumps(metadata.get("image_orientation", []))
                    if metadata.get("image_orientation")
                    else ""
                ),
                image_position=(
                    json.dumps(metadata.get("image_position", []))
                    if metadata.get("image_position")
                    else ""
                ),
                pixel_spacing=(
                    json.dumps(metadata.get("pixel_spacing", []))
                    if metadata.get("pixel_spacing")
                    else ""
                ),
                slice_thickness=metadata.get("slice_thickness"),
                status="Not Assigned",
                assigned_doctors="",
                reported_by="",
                patient_history=request.POST.get("patient_history", ""),
                is_emergency=request.POST.get("is_emergency", "false").lower() == "true",
            )
            logger.info(f"Successfully saved DICOM image to database: ID={dicom_image.id}")
        except Exception as e:
            logger.error(f"Database error: {str(e)}")
            traceback.print_exc()
            try:
                if os.path.exists(full_path):
                    os.remove(full_path)
                    logger.info(f"Cleaned up file after database error: {full_path}")
            except Exception as cleanup_error:
                logger.error(f"Error cleaning up file: {str(cleanup_error)}")
            return JsonResponse(
                {"success": False, "error": f"Database error: {str(e)}"}, 
                status=500
            )
        response_data = {
            "success": True,
            "message": "DICOM file processed successfully",
            "image_id": dicom_image.id,
            "filename": os.path.basename(out_filename),
            "center_name": center_name,
            "file_type": "converted_image" if is_converted_from_image else "dicom",
            "patient_id": metadata.get("patient_id", ""),
            "patient_name": metadata.get("patient_name", ""),
        }
        response = JsonResponse(response_data)
        response["Access-Control-Allow-Origin"] = "*"
        response["Access-Control-Allow-Methods"] = "POST, OPTIONS"
        response["Access-Control-Allow-Headers"] = "Content-Type"
        logger.info(f"Request completed successfully: {response_data}")
        return response
    except Exception as e:
        logger.error(f"Unexpected error in receive_dicom_data: {str(e)}")
        traceback.print_exc()
        return JsonResponse(
            {"success": False, "error": f"Server error: {str(e)}"}, 
            status=500
        )

class PatientPagination:
    def __init__(self, page_size=10):
        self.page_size = page_size





def _get_doctor_group_names(doctor_name):
    all_groups = _read_doctor_groups()
    group_names = []
    for grp in all_groups:
        members = grp.get('members') or []
        if doctor_name in [m.strip() for m in members]:
            group_names.append(grp.get('name', ''))
    return [g for g in group_names if g]


def _build_doctor_study_filter(doctor_name):
    group_names = _get_doctor_group_names(doctor_name)
    q = Q(assigned_doctors__icontains=doctor_name)
    for gname in group_names:
        q |= Q(assigned_doctors__icontains=gname)
    return q


def get_subadmin_allowed_center_names(user):
    cache_key = f"subadmin_allowed_centers_{user.id}"
    cached = cache.get(cache_key)
    if cached is not None:
        return None if cached == "__NONE__" else cached

    profile = UserProfile.objects.filter(user=user).first()
    if not profile:
        return None

    default_center = Center.objects.filter(
        institute_name__iexact='default'
    ).first()
    if default_center and user.username in (default_center.allowed_users or []):
        cache.set(cache_key, "__NONE__", 30)
        return None

    all_centers = Center.objects.exclude(institute_name__iexact='default').prefetch_related('center_names')
    allowed_names = []
    for center in all_centers:
        if user.username in (center.allowed_users or []):
            for cn in center.center_names.all():
                allowed_names.append(cn.name)

    cache.set(cache_key, allowed_names, 30)
    return allowed_names


class DICOMImageViewSet(viewsets.ModelViewSet):
    queryset = DICOMImage.objects.all()
    serializer_class = DICOMImageSerializer
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        queryset = super().get_queryset()
        user = self.request.user
        profile = UserProfile.objects.filter(user=user).first()
        role_name = profile.role.name if profile and profile.role else None
        if role_name == 'SubAdmin':
            allowed = get_subadmin_allowed_center_names(user)
            if allowed is not None:
                queryset = queryset.filter(center_name__in=allowed)

        params = self.request.query_params

        center_name = params.get("center_name")
        if center_name:
            queryset = queryset.filter(center_name=center_name)

        patient_name = params.get("patient_name")
        if patient_name:
            queryset = queryset.filter(patient_name__icontains=patient_name)

        patient_id = params.get("patient_id")
        if patient_id:
            queryset = queryset.filter(patient_id__icontains=patient_id)

        status_filter = params.get("status")
        if status_filter and status_filter != "All":
            queryset = queryset.filter(status=status_filter)

        modality = params.get("modality")
        if modality:
            modalities = [m.strip() for m in modality.split(",") if m.strip()]
            if modalities:
                queryset = queryset.filter(modality__in=modalities)

        emergency = params.get("emergency")
        if emergency and emergency.lower() == "true":
            queryset = queryset.filter(is_emergency=True)

        start_date = params.get("start_date")
        if start_date:
            queryset = queryset.filter(study_date__gte=start_date.replace("-", ""))

        end_date = params.get("end_date")
        if end_date:
            queryset = queryset.filter(study_date__lte=end_date.replace("-", ""))

        return queryset.order_by("-created_at")

    @action(detail=False, methods=["get"])
    def centers(self, request):
        queryset = self.get_queryset()
        institute_map = _build_center_name_to_institute_map()
        center_names = queryset.order_by().values_list("center_name", flat=True).distinct()
        seen = {}
        for cn in center_names:
            if not cn:
                continue
            seen[cn] = institute_map.get(cn, cn)
        centers = sorted(
            [{"center_name": cn, "institute_name": inst} for cn, inst in seen.items()],
            key=lambda c: c["institute_name"]
        )
        return Response({"success": True, "centers": centers})

    def list(self, request):
        queryset = self.filter_queryset(self.get_queryset())
        page = int(request.query_params.get("page", 1))
        page_size = int(request.query_params.get("page_size", 10))

        patient_groups = list(
            queryset.values("patient_id")
            .annotate(latest_created_at=Max("created_at"))
            .order_by("-latest_created_at")
        )

        total_patients = len(patient_groups)
        total_pages = (total_patients + page_size - 1) // page_size if page_size else 1
        start_idx = (page - 1) * page_size
        end_idx = start_idx + page_size

        page_patient_ids = [
            row["patient_id"] for row in patient_groups[start_idx:end_idx]
        ]

        page_queryset = queryset.filter(patient_id__in=page_patient_ids).order_by(
            "-created_at"
        )

        context = self.get_serializer_context()
        context["institute_name_map"] = _build_center_name_to_institute_map()
        serializer = self.get_serializer(page_queryset, many=True, context=context)
        return Response(
            {
                "results": serializer.data,
                "count": total_patients,
                "total_pages": total_pages,
                "current_page": page,
                "patients_on_page": len(page_patient_ids),
            }
        )

    @action(detail=False, methods=["post"], permission_classes=[IsAuthenticated])
    def assign_doctors(self, request):
        if not check_user_permission(request.user, 'assign_doctors'):
            return Response(
                {"success": False, "error": "You don't have permission to assign doctors"},
                status=status.HTTP_403_FORBIDDEN
            )
        
        try:
            image_ids = request.data.get("image_ids", [])
            doctor_names = request.data.get("doctor_names", [])
            if not image_ids or not doctor_names:
                return Response(
                    {"success": False, "error": "image_ids and doctor_names required"},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            images = DICOMImage.objects.filter(id__in=image_ids)
            if len(images) != len(image_ids):
                return Response(
                    {"success": False, "error": "Some IDs not found"},
                    status=status.HTTP_404_NOT_FOUND,
                )

            from .models import CenterName, Center as CenterModel
            from django.contrib.auth.models import User as DjangoUser

            def is_doctor_allowed(doctor_name, allowed_users):
                if doctor_name in allowed_users:
                    return True
                profiles = UserProfile.objects.filter(full_name=doctor_name)
                for profile in profiles:
                    if profile.user.username in allowed_users:
                        return True
                users = DjangoUser.objects.filter(username=doctor_name)
                for u in users:
                    try:
                        if u.userprofile.full_name in allowed_users:
                            return True
                    except Exception:
                        pass
                return False

            all_groups = _read_doctor_groups()

            def resolve_to_group_or_names(names_sent):
                names_set = set(n.strip() for n in names_sent if n.strip())
                for grp in all_groups:
                    members_set = set(m.strip() for m in (grp.get('members') or []) if m.strip())
                    if members_set and members_set == names_set:
                        return [grp['name']]
                    group_name = grp.get('name', '')
                    if group_name in names_set:
                        remaining = names_set - {group_name} - members_set
                        result = [group_name]
                        result.extend(sorted(remaining))
                        return result
                return list(names_sent)

            resolved_names = resolve_to_group_or_names(doctor_names)

            updated_count = 0
            for image in images:
                center_obj = None
                cn_obj = CenterName.objects.filter(name=image.center_name).first()
                if cn_obj:
                    center_obj = cn_obj.center
                else:
                    center_obj = CenterModel.objects.filter(institute_name=image.center_name).first()

                allowed = center_obj.allowed_users if center_obj else []
                has_restriction = isinstance(allowed, list) and len(allowed) > 0

                allowed_for_image = []
                for doctor_name in resolved_names:
                    matched_grp = next((g for g in all_groups if g.get('name') == doctor_name), None)
                    if matched_grp is not None:
                        members = matched_grp.get('members') or []
                        if has_restriction:
                            allowed_members = [m for m in members if is_doctor_allowed(m, allowed)]
                        else:
                            allowed_members = list(members)
                        if allowed_members:
                            if doctor_name not in allowed_for_image:
                                allowed_for_image.append(doctor_name)
                            for m in allowed_members:
                                if m not in allowed_for_image:
                                    allowed_for_image.append(m)
                    else:
                        if has_restriction and not is_doctor_allowed(doctor_name, allowed):
                            continue
                        if doctor_name not in allowed_for_image:
                            allowed_for_image.append(doctor_name)

                if not allowed_for_image:
                    continue

                current_doctors = image.assigned_doctors_list
                for doctor_name in allowed_for_image:
                    if doctor_name not in current_doctors:
                        current_doctors.append(doctor_name)
                image.assigned_doctors = ", ".join(current_doctors)
                if image.status == "Not Assigned":
                    image.status = "Unreported"
                image.save()
                updated_count += 1

            if updated_count == 0:
                debug_info = []
                for image in images:
                    cn_obj = CenterName.objects.filter(name=image.center_name).first()
                    center_obj = cn_obj.center if cn_obj else CenterModel.objects.filter(institute_name=image.center_name).first()
                    allowed = center_obj.allowed_users if center_obj else []
                    debug_info.append(f"center={image.center_name}, allowed={allowed}, doctors_sent={doctor_names}")
                return Response(
                    {"success": False, "error": "No studies could be assigned. Users may not be allowed for these institutes.", "debug": debug_info},
                    status=status.HTTP_403_FORBIDDEN,
                )

            return Response(
                {
                    "success": True,
                    "message": f"Assigned to {updated_count} images",
                    "updated_images": updated_count,
                }
            )
        except Exception as e:
            logger.error(f"Error assigning doctors: {str(e)}")
            return Response(
                {"success": False, "error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    @action(detail=True, methods=["patch"], permission_classes=[IsAuthenticated])
    def update_status(self, request, pk=None):
        try:
            image = self.get_object()
            new_status = request.data.get("status")
            reported_by = request.data.get("reported_by", "")
            report_content = request.data.get("report_content", "")
            if not new_status:
                return Response(
                    {"success": False, "error": "status required"},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            valid_statuses = [
                "Not Assigned",
                "Unreported",
                "Draft",
                "Reviewed",
                "Reported",
            ]
            if new_status not in valid_statuses:
                return Response(
                    {
                        "success": False,
                        "error": f"Invalid status. Valid: {valid_statuses}",
                    },
                    status=status.HTTP_400_BAD_REQUEST,
                )
            if not reported_by:
                try:
                    profile = UserProfile.objects.filter(user=request.user).first()
                    reported_by = (
                        profile.full_name
                        if profile and profile.full_name
                        else request.user.username
                    )
                except:
                    reported_by = request.user.username
            current_report_file = image.report_file.name if image.report_file else None
            image.status = new_status
            image.reported_by = reported_by
            if report_content:
                image.report_content = report_content
            image.save()
            related_images = DICOMImage.objects.filter(
                study_instance_uid=image.study_instance_uid
            ).exclude(id=image.id)
            updated_count = 0
            for related_img in related_images:
                related_img.status = new_status
                related_img.reported_by = reported_by
                if current_report_file:
                    related_img.report_file = current_report_file
                related_img.save()
                updated_count += 1
            return Response(
                {
                    "success": True,
                    "message": "Updated",
                    "image": DICOMImageSerializer(image).data,
                    "updated_images": updated_count + 1,
                    "reported_by": reported_by,
                    "study_instance_uid": image.study_instance_uid or "",
                    "modality": image.modality or "",
                }
            )
        except Exception as e:
            logger.error(f"Error updating status: {str(e)}")
            return Response(
                {"success": False, "error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    @action(detail=False, methods=['get'], url_path='by_doctor', url_name='by_doctor')
    def by_doctor(self, request):
        doctor_name = request.query_params.get('doctor_name')
        if not doctor_name:
            return Response(
                {'success': False, 'error': 'doctor_name required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        images = DICOMImage.objects.filter(
            _build_doctor_study_filter(doctor_name)
        )
        
        try:
            profile = UserProfile.objects.filter(
                Q(full_name=doctor_name) | Q(user__username=doctor_name)
            ).first()
            
            if profile:
                assigned_institutions = profile.assigned_institutions.all()
                if assigned_institutions.exists():
                    center_names = []
                    for institution in assigned_institutions:
                        for center_name in institution.center_names.all():
                            center_names.append(center_name.name)
                    
                    if center_names:
                        institution_images = DICOMImage.objects.filter(
                            center_name__in=center_names
                        )

                        not_assigned = institution_images.filter(
                            status='Not Assigned'
                        ).exclude(assigned_doctors__icontains=doctor_name)

                        batch = []
                        BATCH_SIZE = 500
                        for img in not_assigned.iterator(chunk_size=BATCH_SIZE):
                            current_doctors = img.assigned_doctors_list
                            current_doctors.append(doctor_name)
                            img.assigned_doctors = ', '.join(current_doctors)
                            img.status = 'Unreported'
                            batch.append(img)
                            if len(batch) >= BATCH_SIZE:
                                DICOMImage.objects.bulk_update(batch, ['status', 'assigned_doctors'])
                                batch = []
                        if batch:
                            DICOMImage.objects.bulk_update(batch, ['status', 'assigned_doctors'])

                        images = (images | institution_images).distinct()
        except Exception as e:
            logger.error(f'Error checking institutions: {str(e)}')
        
        images = images.order_by('-created_at')

        page = int(request.query_params.get('page', 1))
        page_size = int(request.query_params.get('page_size', 50))
        total_count = images.count()
        start_idx = (page - 1) * page_size
        end_idx = start_idx + page_size
        page_images = images[start_idx:end_idx]

        context = self.get_serializer_context()
        context['institute_name_map'] = _build_center_name_to_institute_map()
        serializer = self.get_serializer(page_images, many=True, context=context)
        
        return Response({
            'success': True,
            'images': serializer.data,
            'count': total_count
        })

    @action(detail=False, methods=["post"])
    def remove_single_doctor(self, request):
        try:
            image_id = request.data.get("image_id")
            doctor_name = request.data.get("doctor_name")
            dicom_image = self.get_queryset().get(id=image_id)
            assigned_doctors = (
                dicom_image.assigned_doctors.split(",")
                if dicom_image.assigned_doctors
                else []
            )
            assigned_doctors = [d.strip() for d in assigned_doctors]
            if doctor_name in assigned_doctors:
                assigned_doctors.remove(doctor_name)
                dicom_image.assigned_doctors = ",".join(assigned_doctors)
                dicom_image.save()
                return Response({"success": True})
            else:
                return Response(
                    {"success": False, "error": "Doctor not found"}, status=400
                )
        except Exception as e:
            return Response({"success": False, "error": str(e)}, status=500)


def get_studies(request):
    try:
        center_name = request.GET.get("center_name", "")
        queryset = DICOMImage.objects.all()
        if center_name:
            queryset = queryset.filter(center_name=center_name)
        studies = (
            queryset.values(
                "study_instance_uid",
                "patient_name",
                "patient_id",
                "study_date",
                "study_description",
                "modality",
                "center_name",
                "status",
                "assigned_doctors",
                "is_emergency",
            )
            .annotate(image_count=Count("id"))
        )
        return JsonResponse({"success": True, "studies": list(studies)})
    except Exception as e:
        logger.error(f"Error getting studies: {str(e)}")
        return JsonResponse({"success": False, "error": str(e)}, status=500)


@csrf_exempt
def get_studies_grouped(request):
    try:
        user = request.user
        profile = UserProfile.objects.filter(user=user).first()
        role_name = profile.role.name if profile and profile.role else None
        
        center_name = request.GET.get("center_name", "")
        patient_name = request.GET.get("patient_name", "")
        patient_id = request.GET.get("patient_id", "")
        
        queryset = DICOMImage.objects.all()
        
        if role_name == 'SubAdmin':
            allowed = get_subadmin_allowed_center_names(user)
            if allowed is not None:
                queryset = queryset.filter(center_name__in=allowed)
        
        if role_name == 'Doctor':
            doctor_name = profile.full_name or user.username

            cache_key = f"doctor_allowed_centers_{user.id}"
            allowed_center_names = cache.get(cache_key)
            if allowed_center_names is None:
                allowed_center_names = set()
                for center in Center.objects.prefetch_related('center_names').all():
                    if doctor_name in (center.allowed_users or []):
                        allowed_center_names.update(cn.name for cn in center.center_names.all())
                allowed_center_names = list(allowed_center_names)
                cache.set(cache_key, allowed_center_names, 30)

            combined_filter = _build_doctor_study_filter(doctor_name)
            if allowed_center_names:
                combined_filter |= Q(center_name__in=allowed_center_names)

            queryset = queryset.filter(combined_filter).distinct()
        
        if role_name == 'Center':
            center = profile.center if profile else None
            if center:
                center_names = [cn.name for cn in center.center_names.all()]
                queryset = queryset.filter(center_name__in=center_names)
        
        if center_name:
            queryset = queryset.filter(center_name=center_name)
        if patient_name:
            queryset = queryset.filter(patient_name__icontains=patient_name)
        if patient_id:
            queryset = queryset.filter(patient_id__icontains=patient_id)
        studies_dict = {}
        for image in queryset.order_by(
            "study_instance_uid", "series_number", "instance_number"
        ):
            study_uid = image.study_instance_uid
            if study_uid not in studies_dict:
                studies_dict[study_uid] = {
                    "study_instance_uid": study_uid,
                    "patient_name": image.patient_name,
                    "patient_id": image.patient_id,
                    "patient_birth_date": image.patient_birth_date,
                    "patient_sex": image.patient_sex,
                    "study_date": image.study_date,
                    "study_description": image.study_description,
                    "modality": image.modality,
                    "center_name": image.center_name,
                    "status": image.status,
                    "assigned_doctors": image.assigned_doctors,
                    "is_emergency": image.is_emergency,
                    "images": [],
                    "image_count": 0,
                }
            studies_dict[study_uid]["images"].append(
                {
                    "id": image.id,
                    "sop_instance_uid": image.sop_instance_uid,
                    "instance_number": image.instance_number,
                    "series_description": image.series_description,
                    "file_path": image.file_path,
                    "file_size": image.file_size,
                }
            )
            studies_dict[study_uid]["image_count"] += 1
        studies_list = list(studies_dict.values())
        return JsonResponse({"success": True, "studies": studies_list})
    except Exception as e:
        logger.error(f"Error getting grouped studies: {str(e)}")
        return JsonResponse({"success": False, "error": str(e)}, status=500)


def get_centers(request):
    try:
        centers = DICOMImage.get_centers()
        center_list = []
        for center in centers:
            center_stats = DICOMImage.get_center_stats(center["center_name"])
            center_list.append(center_stats)
        return JsonResponse({"success": True, "centers": center_list})
    except Exception as e:
        logger.error(f"Error getting centers: {str(e)}")
        return JsonResponse({"success": False, "error": str(e)}, status=500)


@api_view(["GET"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def get_my_centers(request):
    try:
        user = request.user
        profile = UserProfile.objects.filter(user=user).first()
        role_name = profile.role.name if profile and profile.role else None

        if role_name == 'SubAdmin':
            allowed = get_subadmin_allowed_center_names(user)
            if allowed is not None:
                center_names = allowed
            else:
                center_names = list(
                    DICOMImage.objects.values_list('center_name', flat=True).distinct()
                )
        else:
            center_names = list(
                DICOMImage.objects.values_list('center_name', flat=True).distinct()
            )

        institute_map = _build_center_name_to_institute_map()
        centers = [
            {
                'center_name': name,
                'institute_name': institute_map.get(name, name),
            }
            for name in sorted(set(c for c in center_names if c))
        ]
        return Response({"success": True, "centers": centers})
    except Exception as e:
        return Response({"success": False, "error": str(e)}, status=500)


def get_center_detail(request, center_name):
    try:
        center_stats = DICOMImage.get_center_stats(center_name)
        recent_images = DICOMImage.objects.filter(center_name=center_name).order_by(
            "-created_at"
        )[:10]
        images_data = []
        for image in recent_images:
            images_data.append(
                {
                    "id": image.id,
                    "patient_name": image.patient_name,
                    "patient_id": image.patient_id,
                    "study_date": image.study_date,
                    "study_description": image.study_description,
                    "modality": image.modality,
                    "file_size_mb": image.file_size_mb,
                    "status": image.status,
                    "assigned_doctors": image.assigned_doctors_list,
                    "created_at": (
                        image.created_at.isoformat() if image.created_at else None
                    ),
                }
            )
        return JsonResponse(
            {
                "success": True,
                "center_stats": center_stats,
                "recent_images": images_data,
            }
        )
    except Exception as e:
        logger.error(f"Error getting center detail: {str(e)}")
        return JsonResponse({"success": False, "error": str(e)}, status=500)


def get_study_detail(request, study_id):
    try:
        images = DICOMImage.objects.filter(id=study_id)
        if not images.exists():
            return JsonResponse(
                {"success": False, "error": "Study not found"}, status=404
            )
        study_data = []
        for image in images:
            study_data.append(
                {
                    "id": image.id,
                    "center_name": image.center_name,
                    "patient_name": image.patient_name,
                    "patient_id": image.patient_id,
                    "study_date": image.study_date,
                    "study_description": image.study_description,
                    "series_description": image.series_description,
                    "modality": image.modality,
                    "instance_number": image.instance_number,
                    "file_size": image.file_size,
                    "status": image.status,
                    "assigned_doctors": image.assigned_doctors_list,
                    "created_at": image.created_at,
                }
            )
        return JsonResponse({"success": True, "study": study_data})
    except Exception as e:
        logger.error(f"Error getting study detail: {str(e)}")
        return JsonResponse({"success": False, "error": str(e)}, status=500)


def get_study_images(request, study_uid):
    try:
        images = DICOMImage.objects.filter(study_instance_uid=study_uid).order_by(
            "series_number", "instance_number"
        )
        image_list = []
        for image in images:
            image_list.append(
                {
                    "id": image.id,
                    "center_name": image.center_name,
                    "sop_instance_uid": image.sop_instance_uid,
                    "instance_number": image.instance_number,
                    "series_description": image.series_description,
                    "file_path": image.file_path,
                    "file_size": image.file_size,
                    "status": image.status,
                    "assigned_doctors": image.assigned_doctors_list,
                    "patient_name": image.patient_name,
                    "patient_id": image.patient_id,
                    "study_date": image.study_date,
                    "modality": image.modality,
                    "series_number": image.series_number,
                    "series_instance_uid": image.series_instance_uid,
                }
            )
        return JsonResponse({"success": True, "images": image_list})
    except Exception as e:
        logger.error(f"Error getting study images: {str(e)}")
        return JsonResponse({"success": False, "error": str(e)}, status=500)


def get_stats(request):
    try:
        center_name = request.GET.get("center_name", "")
        if center_name:
            stats = DICOMImage.get_center_stats(center_name)
        else:
            total_images = DICOMImage.objects.count()
            total_studies = (
                DICOMImage.objects.values("study_instance_uid").distinct().count()
            )
            total_patients = DICOMImage.objects.values("patient_id").distinct().count()
            total_centers = DICOMImage.objects.values("center_name").distinct().count()
            status_counts = DICOMImage.objects.values("status").annotate(
                count=Count("id")
            )
            total_size = (
                DICOMImage.objects.aggregate(total=Sum("file_size"))["total"] or 0
            )
            stats = {
                "total_images": total_images,
                "total_studies": total_studies,
                "total_patients": total_patients,
                "total_centers": total_centers,
                "total_size_bytes": total_size,
                "total_size_mb": (
                    round(total_size / (1024 * 1024), 2) if total_size else 0
                ),
                "status_breakdown": {
                    item["status"]: item["count"] for item in status_counts
                },
            }
        return JsonResponse({"success": True, "stats": stats})
    except Exception as e:
        logger.error(f"Error getting stats: {str(e)}")
        return JsonResponse({"success": False, "error": str(e)}, status=500)


def api_info(request):
    return JsonResponse(
        {"success": True, "message": "Telerad PACS API", "version": "2.0"}
    )


def test_api(request):
    return JsonResponse(
        {
            "success": True,
            "message": "API working",
            "timestamp": datetime.now().isoformat(),
        }
    )


def get_all_dicom_images(request):
    try:
        center_name = request.GET.get("center_name", "")
        doctor_name = request.GET.get("doctor_name", "")
        status_filter = request.GET.get("status", "")
        queryset = DICOMImage.objects.all().order_by("-created_at")
        if center_name:
            queryset = queryset.filter(center_name=center_name)
        if doctor_name:
            queryset = queryset.filter(assigned_doctors__icontains=doctor_name)
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        images_list = []
        for image in queryset:
            images_list.append(
                {
                    "id": image.id,
                    "center_name": image.center_name,
                    "patient_name": image.patient_name,
                    "patient_id": image.patient_id,
                    "study_date": image.study_date,
                    "study_description": image.study_description,
                    "series_description": image.series_description,
                    "modality": image.modality,
                    "study_instance_uid": image.study_instance_uid,
                    "instance_number": image.instance_number,
                    "file_size": image.file_size,
                    "file_path": image.file_path,
                    "status": image.status,
                    "assigned_doctors": image.assigned_doctors,
                    "assigned_doctors_list": image.assigned_doctors_list,
                    "reported_by": image.reported_by,
                    "is_emergency": image.is_emergency,
                    "created_at": (
                        image.created_at.isoformat() if image.created_at else None
                    ),
                }
            )
        return JsonResponse({"success": True, "images": images_list})
    except Exception as e:
        logger.error(f"Error getting images: {str(e)}")
        return JsonResponse({"success": False, "error": str(e)}, status=500)


def get_fixed_stats(request):
    try:
        center_name = request.GET.get("center_name", "")
        if center_name:
            queryset = DICOMImage.objects.filter(center_name=center_name)
        else:
            queryset = DICOMImage.objects.all()
        patient_ids = []
        total_size = 0
        for image in queryset:
            if image.patient_id and image.patient_id not in patient_ids:
                patient_ids.append(image.patient_id)
            if image.file_size:
                total_size += image.file_size
        stats = {
            "patients": len(patient_ids),
            "size_bytes": total_size,
            "size_mb": round(total_size / (1024 * 1024), 2) if total_size else 0,
        }
        if center_name:
            stats["center_name"] = center_name
        return JsonResponse({"success": True, "stats": stats})
    except Exception as e:
        logger.error(f"Error getting stats: {str(e)}")
        return JsonResponse({"success": False, "error": str(e)}, status=500)


def index(request):
    return render(request, "/static/index.html")


@csrf_exempt
def assign_doctors_to_images(request):
    if request.method != "POST":
        return JsonResponse(
            {"success": False, "error": "Only POST allowed"}, status=405
        )
    try:
        data = json.loads(request.body)
        image_ids = data.get("image_ids", [])
        doctor_names = data.get("doctor_names", [])
        
        if not image_ids:
            return JsonResponse(
                {"success": False, "error": "image_ids required"},
                status=400,
            )
        
        if not doctor_names or len(doctor_names) == 0:
            return JsonResponse(
                {"success": False, "error": "doctor_names required"},
                status=400,
            )
        
        images = DICOMImage.objects.filter(id__in=image_ids)
        if len(images) != len(image_ids):
            return JsonResponse(
                {"success": False, "error": "Some IDs not found"}, status=404
            )
        all_groups = _read_doctor_groups()

        def resolve_to_group_or_names(names_sent):
            names_set = set(n.strip() for n in names_sent if n.strip())
            for grp in all_groups:
                members_set = set(m.strip() for m in (grp.get('members') or []) if m.strip())
                if members_set and members_set == names_set:
                    return [grp['name']]
                group_name = grp.get('name', '')
                if group_name in names_set:
                    remaining = names_set - {group_name} - members_set
                    result = [group_name]
                    result.extend(sorted(remaining))
                    return result
            return list(names_sent)

        resolved_names = resolve_to_group_or_names(doctor_names)

        def is_doctor_allowed_simple(doctor_name, allowed_users):
            if doctor_name in allowed_users:
                return True
            profiles = UserProfile.objects.filter(full_name=doctor_name)
            for profile in profiles:
                if profile.user.username in allowed_users:
                    return True
            return False

        updated_count = 0
        for image in images:
            cn_obj = CenterName.objects.filter(name=image.center_name).first()
            if cn_obj:
                center_obj = cn_obj.center
            else:
                center_obj = Center.objects.filter(institute_name=image.center_name).first()
            allowed = center_obj.allowed_users if center_obj else []
            has_restriction = isinstance(allowed, list) and len(allowed) > 0

            names_to_assign = []
            for doctor_name in resolved_names:
                matched_grp = next((g for g in all_groups if g.get('name') == doctor_name), None)
                if matched_grp is not None:
                    members = matched_grp.get('members') or []
                    if has_restriction:
                        allowed_members = [m for m in members if is_doctor_allowed_simple(m, allowed)]
                    else:
                        allowed_members = list(members)
                    if allowed_members:
                        if doctor_name not in names_to_assign:
                            names_to_assign.append(doctor_name)
                        for m in allowed_members:
                            if m not in names_to_assign:
                                names_to_assign.append(m)
                else:
                    if has_restriction and not is_doctor_allowed_simple(doctor_name, allowed):
                        continue
                    if doctor_name not in names_to_assign:
                        names_to_assign.append(doctor_name)

            for name in names_to_assign:
                image.assign_doctor(name)
            updated_count += 1
        return JsonResponse(
            {
                "success": True,
                "message": f"Assigned to {updated_count} images",
                "updated_images": updated_count,
            }
        )
    except json.JSONDecodeError:
        return JsonResponse({"success": False, "error": "Invalid JSON"}, status=400)
    except Exception as e:
        logger.error(f"Error in assign_doctors_to_images: {str(e)}")
        return JsonResponse({"success": False, "error": str(e)}, status=500)


@method_decorator(csrf_exempt, name="dispatch")
class ReceiveDICOMView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        return receive_dicom_data(request)


class DICOMImageFilter(filters.FilterSet):
    patient_name__icontains = filters.CharFilter(
        field_name="patient_name", lookup_expr="icontains"
    )
    patient_id__icontains = filters.CharFilter(
        field_name="patient_id", lookup_expr="icontains"
    )
    status = filters.CharFilter(field_name="status")
    center_name = filters.CharFilter(field_name="center_name")
    is_emergency = filters.BooleanFilter(field_name="is_emergency")
    modality__in = filters.CharFilter(method="filter_modality")
    date_from = filters.CharFilter(method="filter_date_from")
    date_to = filters.CharFilter(method="filter_date_to")
    institute_name = filters.CharFilter(method="filter_institute_name")

    class Meta:
        model = DICOMImage
        fields = [
            "patient_name__icontains",
            "patient_id__icontains",
            "status",
            "center_name",
            "is_emergency",
        ]

    def filter_modality(self, queryset, name, value):
        if value:
            modalities = [m.strip() for m in value.split(",") if m.strip()]
            return queryset.filter(modality__in=modalities)
        return queryset

    def filter_date_from(self, queryset, name, value):
        if value:
            return queryset.filter(study_date__gte=value)
        return queryset

    def filter_date_to(self, queryset, name, value):
        if value:
            return queryset.filter(study_date__lte=value)
        return queryset

    def filter_institute_name(self, queryset, name, value):
        if value:
            from .models import CenterName
            center_names = CenterName.objects.filter(
                center__institute_name__iexact=value
            ).values_list('name', flat=True)
            return queryset.filter(center_name__in=list(center_names))
        return queryset


class StudyGroupedPagination(PageNumberPagination):
    page_size = 100
    page_size_query_param = 'page_size'
    max_page_size = 1000


class DICOMImageListView(generics.ListAPIView):
    serializer_class = DICOMImageSerializer
    pagination_class = StudyGroupedPagination
    filterset_class = DICOMImageFilter
    filter_backends = [filters.DjangoFilterBackend]
    authentication_classes = [TokenAuthentication]
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        profile = UserProfile.objects.filter(user=user).first()
        role_name = profile.role.name if profile and profile.role else None
        
        if role_name == 'SubAdmin':
            allowed = get_subadmin_allowed_center_names(user)
            if allowed is not None:
                return DICOMImage.objects.filter(center_name__in=allowed).order_by("-created_at")
        
        if role_name == 'Doctor':
            doctor_name = profile.full_name or user.username

            cache_key = f"doctor_allowed_centers_{user.id}"
            allowed_center_names = cache.get(cache_key)
            if allowed_center_names is None:
                allowed_center_names = set()
                for center in Center.objects.prefetch_related('center_names').all():
                    if doctor_name in (center.allowed_users or []):
                        allowed_center_names.update(cn.name for cn in center.center_names.all())
                allowed_center_names = list(allowed_center_names)
                cache.set(cache_key, allowed_center_names, 30)

            combined_filter = _build_doctor_study_filter(doctor_name)
            if allowed_center_names:
                combined_filter |= Q(center_name__in=allowed_center_names)

            return DICOMImage.objects.filter(combined_filter).order_by("-created_at").distinct()
        
        if role_name == 'Center':
            center = profile.center if profile else None
            if center:
                center_names = [cn.name for cn in center.center_names.all()]
                return DICOMImage.objects.filter(center_name__in=center_names).order_by("-created_at")
        
        return DICOMImage.objects.all().order_by("-created_at")

    def list(self, request, *args, **kwargs):
        from django.db.models import OuterRef, Subquery

        queryset = self.filter_queryset(self.get_queryset())

        latest_id_for_uid = Subquery(
            DICOMImage.objects.filter(
                study_instance_uid=OuterRef('study_instance_uid')
            ).order_by('-id').values('id')[:1]
        )

        deduped_qs = queryset.filter(
            Q(study_instance_uid__isnull=True) |
            Q(study_instance_uid__exact='') |
            Q(id=latest_id_for_uid)
        ).order_by('-created_at')

        page_size = max(1, min(int(request.query_params.get('page_size', 100)), 500))
        page_num = max(1, int(request.query_params.get('page', 1)))
        start = (page_num - 1) * page_size
        end = start + page_size

        total = deduped_qs.count()
        page_records = deduped_qs[start:end]

        context = self.get_serializer_context()
        context['institute_name_map'] = _build_center_name_to_institute_map()
        serializer = self.get_serializer(page_records, many=True, context=context)
        next_url = None
        if end < total:
            next_url = "has_more"

        return Response({
            'count': total,
            'next': next_url,
            'previous': None,
            'results': serializer.data,
        })

@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([AllowAny])
def current_user(request):
    user = request.user
    if not user or not user.is_authenticated:
        return Response(
            {'success': False, 'error': 'Not authenticated'},
            status=status.HTTP_401_UNAUTHORIZED
        )
    
    try:
        profile = UserProfile.objects.filter(user=user).first()
        
        if profile:
            doctor_name = profile.full_name if profile.full_name else user.username
            role = profile.role.name if profile.role else 'Unknown'
        else:
            doctor_name = user.username
            role = 'Admin' if user.is_superuser else 'Unknown'
        
        center_name = None
        institute_name = None
        assigned_institutions = []
        
        can_assign_doctors = False
        can_write_reports = False
        can_manage_templates = False
        can_view_images = False
        
        if user.is_superuser:
            can_assign_doctors = True
            can_write_reports = True
            can_manage_templates = True
            can_view_images = True
        elif profile:
            can_assign_doctors = profile.can_assign_doctors
            can_write_reports = profile.can_write_reports
            can_manage_templates = profile.can_manage_templates
            can_view_images = profile.can_view_images
        
        if role == 'Doctor' and profile:
            assigned_institutions = [
                {
                    'id': inst.id,
                    'name': inst.institute_name,
                    'centers': [cn.name for cn in inst.center_names.all()]
                }
                for inst in profile.assigned_institutions.all()
            ]
        
        if role == 'Center':
            if profile and profile.center:
                center_obj = profile.center
                institute_name = center_obj.institute_name
                center_name_obj = center_obj.center_names.first()
                center_name = center_name_obj.name if center_name_obj else None
            else:
                center = Center.objects.filter(user=user).first()
                if center:
                    institute_name = center.institute_name
                    center_name_obj = center.center_names.first()
                    center_name = center_name_obj.name if center_name_obj else None
                    if profile:
                        profile.center = center
                        profile.save()
        
        response_data = {
            'success': True,
            'username': user.username,
            'doctor_name': doctor_name,
            'full_name': profile.full_name if profile else '',
            'designation': profile.designation if profile else '',
            'qualification': profile.qualification if profile else '',
            'contact_number': profile.contact_number if profile else '',
            'bmdc_reg_no': profile.bmdc_reg_no if profile else '',
            'signature': profile.signature.url if (profile and profile.signature) else '',
            'role': role,
            'center_name': center_name,
            'institute_name': institute_name,
            'assigned_institutions': assigned_institutions,
            'permissions': {
                'can_assign_doctors': can_assign_doctors,
                'can_write_reports': can_write_reports,
                'can_manage_templates': can_manage_templates,
                'can_view_images': can_view_images
            }
        }
        
        return Response(response_data, status=status.HTTP_200_OK)
        
    except Exception as e:
        logger.error(f'ERROR in current_user: {str(e)}')
        return Response(
            {'success': False, 'error': f'Server error: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

@api_view(["GET"])
@permission_classes([IsAuthenticated])
def get_doctors(request):
    try:
        cache_key = "doctors_list_v1"
        doctors = cache.get(cache_key)
        if doctors is None:
            doctor_profiles = (
                UserProfile.objects.filter(role__name="Doctor")
                .select_related("user")
                .values(
                    "user__id",
                    "user__username",
                    "full_name",
                    "designation",
                    "qualification",
                    "bmdc_reg_no",
                )
            )
            doctors = [
                {
                    "id": profile["user__id"],
                    "name": profile["full_name"] or profile["user__username"],
                    "username": profile["user__username"],
                    "designation": profile["designation"],
                    "qualification": profile["qualification"],
                    "bmdc_reg_no": profile["bmdc_reg_no"],
                }
                for profile in doctor_profiles
            ]
            cache.set(cache_key, doctors, 30)
        return Response({"success": True, "doctors": doctors})
    except Exception as e:
        return Response({"success": False, "error": str(e)}, status=500)


@api_view(["GET"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def get_current_user_info(request):
    try:
        user = request.user
        if not user.is_authenticated:
            return Response(
                {"success": False, "error": "Not authenticated"},
                status=status.HTTP_401_UNAUTHORIZED,
            )
        profile = UserProfile.objects.filter(user=user).first()
        if not profile:
            return Response(
                {"success": False, "error": "Profile not found"},
                status=status.HTTP_404_NOT_FOUND,
            )
        role_name = profile.role.name if profile.role else "Unknown"
        if role_name not in ["SubAdmin", "Center"]:
            return Response(
                {"success": False, "error": "Access denied"},
                status=status.HTTP_403_FORBIDDEN,
            )
        user_name = profile.full_name or user.username
        center_name = (
            profile.center.name if role_name == "Center" and profile.center else None
        )
        return Response(
            {
                "success": True,
                "username": user.username,
                "display_name": user_name,
                "role": role_name,
                "center_name": center_name,
            }
        )
    except Exception as e:
        logger.error(f"Error in get_current_user_info: {str(e)}")
        return Response(
            {"success": False, "error": str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

@api_view(["POST"])
@permission_classes([IsAuthenticated])
def upload_dicom_report(request, dicom_id):
    if not check_user_permission(request.user, 'write_reports'):
        return Response(
            {"success": False, "error": "You don't have permission to write reports"},
            status=status.HTTP_403_FORBIDDEN
        )
    
    try:
        dicom = DICOMImage.objects.get(id=dicom_id)
        if "file" not in request.FILES:
            return Response({"success": False, "error": "No file provided"}, status=400)

        # ── Rescue orphaned old report BEFORE overwriting reported_by ─────────
        # This runs first, so dicom.reported_by is still the OLD doctor's name.
        # We only rescue if there is no history file yet (first save with new code).
        try:
            if dicom.report_content and dicom.report_content.strip() and dicom.reported_by:
                existing = _read_history(dicom_id)
                if not existing:
                    old_file_url = None
                    if dicom.report_file:
                        try:
                            old_file_url = request.build_absolute_uri(dicom.report_file.url)
                        except Exception:
                            pass
                    dt_str = ''
                    if dicom.updated_at:
                        dt_str = dicom.updated_at.strftime('%d-%m-%Y %H:%M:%S')
                    _append_to_history(dicom_id, {
                        'reported_by': dicom.reported_by,
                        'report_content': dicom.report_content,
                        'status': dicom.status or 'Reported',
                        'report_file': old_file_url,
                        'created_at': dt_str,
                        'dicom_id': dicom_id,
                    })
        except Exception as rescue_err:
            logger.warning(f"Could not rescue old report for {dicom_id}: {rescue_err}")
        # ─────────────────────────────────────────────────────────────────────

        file = request.FILES["file"]
        clean_filename = file.name.split("/")[-1].split("\\")[-1]
        final_filename = f"{dicom.patient_id}_{dicom_id}_{clean_filename}"
        dicom.report_file.save(final_filename, file, save=False)
        try:
            profile = UserProfile.objects.filter(user=request.user).first()
            doctor_name = (
                profile.full_name
                if profile and profile.full_name
                else request.user.username
            )
        except:
            doctor_name = request.user.username
        dicom.status = "Reported"
        dicom.reported_by = doctor_name
        dicom.save()
        saved_report_path = dicom.report_file.name

        try:
            report_file_url = None
            if dicom.report_file:
                try:
                    report_file_url = request.build_absolute_uri(dicom.report_file.url)
                except Exception:
                    pass
            from datetime import datetime
            _append_to_history(dicom_id, {
                'reported_by': doctor_name,
                'report_content': dicom.report_content or '',
                'status': 'Reported',
                'report_file': report_file_url,
                'created_at': datetime.now().strftime('%d-%m-%Y %H:%M:%S'),
                'dicom_id': dicom_id,
            })
        except Exception as hist_err:
            logger.warning(f"Could not save report history for {dicom_id}: {hist_err}")

        return Response(
            {
                "success": True,
                "message": "Report uploaded",
                "file_path": saved_report_path,
                "updated_images": 1,
                "reported_by": doctor_name,
            }
        )
    except DICOMImage.DoesNotExist:
        return Response({"success": False, "error": "DICOM not found"}, status=404)
    except Exception as e:
        logger.error(f"Upload error: {str(e)}")
        return Response({"success": False, "error": str(e)}, status=500)


@api_view(["PATCH"])
@permission_classes([IsAuthenticated])
def update_dicom_status(request, dicom_id):
    if not check_user_permission(request.user, 'write_reports'):
        return Response(
            {"success": False, "error": "You don't have permission to update report status"},
            status=status.HTTP_403_FORBIDDEN
        )
    try:
        dicom = DICOMImage.objects.get(id=dicom_id)
        new_status = request.data.get("status")
        reported_by = request.data.get("reported_by", "")
        report_content = request.data.get("report_content", "")
        if not new_status:
            return Response(
                {"success": False, "error": "Status required"},
                status=status.HTTP_400_BAD_REQUEST,
            )
        valid_statuses = ["Not Assigned", "Unreported", "Draft", "Reviewed", "Reported"]
        if new_status not in valid_statuses:
            return Response(
                {"success": False, "error": f"Invalid status. Valid: {valid_statuses}"},
                status=status.HTTP_400_BAD_REQUEST,
            )
        if not reported_by:
            try:
                profile = UserProfile.objects.filter(user=request.user).first()
                reported_by = (
                    profile.full_name
                    if profile and profile.full_name
                    else request.user.username
                )
            except:
                reported_by = request.user.username
        current_report_file = dicom.report_file.name if dicom.report_file else None

        dicom.status = new_status
        dicom.reported_by = reported_by
        if report_content:
            dicom.report_content = report_content
        dicom.save()

        # ── Append new entry to file-based history (no migration needed) ──────
        if report_content:
            try:
                report_file_url = None
                if dicom.report_file:
                    try:
                        report_file_url = request.build_absolute_uri(dicom.report_file.url)
                    except Exception:
                        pass
                from datetime import datetime
                _append_to_history(dicom_id, {
                    'reported_by': reported_by,
                    'report_content': report_content,
                    'status': new_status,
                    'report_file': report_file_url,
                    'created_at': datetime.now().strftime('%d-%m-%Y %H:%M:%S'),
                    'dicom_id': dicom_id,
                })
            except Exception as hist_err:
                logger.warning(f"Could not save report history for {dicom_id}: {hist_err}")
        # ──────────────────────────────────────────────────────────────────────

        related_images = DICOMImage.objects.filter(
            study_instance_uid=dicom.study_instance_uid
        ).exclude(id=dicom_id)
        updated_count = 0
        for related_img in related_images:
            related_img.status = new_status
            related_img.reported_by = reported_by
            if current_report_file:
                related_img.report_file = current_report_file
            related_img.save()
            updated_count += 1
        return Response(
            {
                "success": True,
                "message": "Updated",
                "image": DICOMImageSerializer(dicom).data,
                "updated_images": updated_count + 1,
                "reported_by": reported_by,
                "study_instance_uid": dicom.study_instance_uid or "",
                "modality": dicom.modality or "",
            }
        )
    except DICOMImage.DoesNotExist:
        return Response(
            {"success": False, "error": "DICOM not found"},
            status=status.HTTP_404_NOT_FOUND,
        )
    except Exception as e:
        logger.error(f"Error updating status: {str(e)}")
        return Response(
            {"success": False, "error": str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def get_report_history(request, dicom_id):
    try:
        dicom = DICOMImage.objects.get(id=dicom_id)

        # Read from file-based history (no database table needed)
        result = _read_history(dicom_id)

        # Fallback: if no history file yet, return the single report on DICOMImage
        if not result and dicom.report_content and dicom.report_content.strip():
            report_file_url = None
            if dicom.report_file:
                try:
                    report_file_url = request.build_absolute_uri(dicom.report_file.url)
                except Exception:
                    pass
            dt_str = ''
            if dicom.updated_at:
                dt_str = dicom.updated_at.strftime('%d-%m-%Y %H:%M:%S')
            result = [{
                'reported_by': dicom.reported_by or '',
                'status': dicom.status or '',
                'report_file': report_file_url,
                'report_content': dicom.report_content,
                'created_at': dt_str,
                'dicom_id': dicom.id,
            }]

        return Response({"success": True, "history": result, "count": len(result)})
    except DICOMImage.DoesNotExist:
        return Response({"success": False, "error": "DICOM not found"}, status=404)
    except Exception as e:
        return Response({"success": False, "error": str(e)}, status=500)


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def get_patient_history_data(request):
    patient_id = request.query_params.get("patient_id", "").strip()
    if not patient_id:
        return Response({"success": False, "error": "patient_id required"}, status=400)

    try:
        images = DICOMImage.objects.filter(patient_id=patient_id).order_by("-id")

        seen_uids = set()
        unique = []
        for img in images:
            uid = img.study_instance_uid or str(img.id)
            if uid not in seen_uids:
                seen_uids.add(uid)
                unique.append(img)

        if not unique:
            return Response({"success": False, "error": "No studies found"}, status=404)

        first = unique[0]

        age = 0
        if first.patient_birth_date:
            try:
                from datetime import datetime as dt
                bd = dt.strptime(first.patient_birth_date, "%Y%m%d")
                today = dt.now()
                age = today.year - bd.year
                if (today.month, today.day) < (bd.month, bd.day):
                    age -= 1
                age = max(0, age)
            except Exception:
                age = 0

        def fmt_date(study_date, study_time):
            if not study_date:
                return ""
            try:
                y = str(study_date)[:4]
                m = str(study_date)[4:6]
                d = str(study_date)[6:8]
                if study_time:
                    t = str(study_time)
                    h, mn, s = t[:2], t[2:4], t[4:6] if len(t) >= 6 else "00"
                    return f"{d}-{m}-{y} {h}:{mn}:{s}"
                return f"{d}-{m}-{y}"
            except Exception:
                return str(study_date)

        def get_institute(img):
            try:
                cn = CenterName.objects.filter(name=img.center_name).select_related("center").first()
                if cn and cn.center and cn.center.institute_name:
                    return cn.center.institute_name
                center = Center.objects.filter(center_names__name=img.center_name).first()
                if center and center.institute_name:
                    return center.institute_name
            except Exception:
                pass
            return img.center_name

        institute_name = get_institute(first)

        patient_info = {
            "patient_name": first.patient_name,
            "patient_id": first.patient_id,
            "patient_sex": first.patient_sex,
            "age": age,
            "modality": first.modality,
            "institute_name": institute_name,
            "study_date_display": fmt_date(first.study_date, first.study_time),
        }

        studies = []
        for img in unique:
            referral_file_url = None
            if img.referral_file:
                try:
                    referral_file_url = request.build_absolute_uri(img.referral_file.url)
                except Exception:
                    pass

            studies.append({
                "id": img.id,
                "study_description": img.study_description or "",
                "referring_physician": img.referring_physician or "",
                "is_emergency": img.is_emergency,
                "patient_history": img.patient_history or "",
                "referral_file": referral_file_url,
                "study_date_display": fmt_date(img.study_date, img.study_time),
            })

        return Response({"success": True, "patient": patient_info, "studies": studies})

    except Exception as e:
        return Response({"success": False, "error": str(e)}, status=500)


@api_view(["POST"])
@permission_classes([IsAuthenticated])
def remove_single_doctor(request):
    try:
        image_id = request.data.get("image_id")
        doctor_name = request.data.get("doctor_name")
        dicom_image = DICOMImage.objects.get(id=image_id)
        assigned_doctors = list(dicom_image.assigned_doctors_list or [])
        if doctor_name in assigned_doctors:
            assigned_doctors.remove(doctor_name)
            dicom_image.assigned_doctors_list = assigned_doctors
            dicom_image.save()
            return Response({"success": True})
        else:
            return Response({"success": False, "error": "Doctor not found"}, status=400)
    except DICOMImage.DoesNotExist:
        return Response({"success": False, "error": "Image not found"}, status=404)
    except Exception as e:
        return Response({"success": False, "error": str(e)}, status=500)


@api_view(["GET"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def get_institute_info(request):
    try:
        user = request.user
        if not user.is_authenticated:
            return Response(
                {"success": False, "error": "Not authenticated"},
                status=status.HTTP_401_UNAUTHORIZED,
            )
        profile = UserProfile.objects.filter(user=user).first()
        if not profile:
            return Response(
                {"success": False, "error": "Profile not found"},
                status=status.HTTP_404_NOT_FOUND,
            )
        role_name = profile.role.name if profile.role else "Unknown"
        if role_name != "Center":
            return Response(
                {"success": False, "error": "Access denied"},
                status=status.HTTP_403_FORBIDDEN,
            )
        if profile.center:
            institute_name = profile.center.institute_name
            centers_in_institute = Center.objects.filter(institute_name=institute_name)
            center_names_list = []
            for center_obj in centers_in_institute:
                for center_name in center_obj.center_names.all():
                    center_names_list.append(
                        {"name": center_name.name, "id": center_name.id}
                    )
            return Response(
                {
                    "success": True,
                    "username": user.username,
                    "role": role_name,
                    "institute_name": institute_name,
                    "centers": center_names_list,
                    "center_count": len(center_names_list),
                }
            )
        else:
            return Response(
                {"success": False, "error": "No institute assigned"},
                status=status.HTTP_404_NOT_FOUND,
            )
    except Exception as e:
        logger.error(f"Error in get_institute_info: {str(e)}")
        return Response(
            {"success": False, "error": str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@api_view(["GET"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def get_institute_studies(request):
    try:
        user = request.user
        profile = UserProfile.objects.filter(user=user).first()
        if not profile or not profile.center:
            return Response(
                {"success": False, "error": "No institute assigned"},
                status=status.HTTP_404_NOT_FOUND,
            )
        institute_name = profile.center.institute_name
        centers_in_institute = Center.objects.filter(institute_name=institute_name)
        center_names_list = []
        for center_obj in centers_in_institute:
            for center_name in center_obj.center_names.all():
                center_names_list.append(center_name.name)

        queryset = DICOMImage.objects.filter(center_name__in=center_names_list)

        center_filter = request.GET.get("center_name")
        if center_filter:
            queryset = queryset.filter(center_name=center_filter)

        status_filter = request.GET.get("status")
        if status_filter and status_filter != "All":
            queryset = queryset.filter(status=status_filter)

        patient_name = request.GET.get("patient_name")
        if patient_name:
            queryset = queryset.filter(patient_name__icontains=patient_name)

        patient_id = request.GET.get("patient_id")
        if patient_id:
            queryset = queryset.filter(patient_id__icontains=patient_id)

        modality = request.GET.get("modality")
        if modality:
            modalities = [m.strip() for m in modality.split(",") if m.strip()]
            if modalities:
                queryset = queryset.filter(modality__in=modalities)

        emergency = request.GET.get("emergency")
        if emergency and emergency.lower() == "true":
            queryset = queryset.filter(is_emergency=True)

        start_date = request.GET.get("start_date")
        if start_date:
            queryset = queryset.filter(study_date__gte=start_date.replace("-", ""))

        end_date = request.GET.get("end_date")
        if end_date:
            queryset = queryset.filter(study_date__lte=end_date.replace("-", ""))

        page = int(request.GET.get("page", 1))
        page_size = int(request.GET.get("page_size", 10))

        patient_groups = list(
            queryset.values("patient_id")
            .annotate(latest_created_at=Max("created_at"))
            .order_by("-latest_created_at")
        )
        total_patients = len(patient_groups)
        total_pages = (total_patients + page_size - 1) // page_size if page_size else 1
        start_idx = (page - 1) * page_size
        end_idx = start_idx + page_size
        page_patient_ids = [row["patient_id"] for row in patient_groups[start_idx:end_idx]]

        page_queryset = queryset.filter(patient_id__in=page_patient_ids).order_by(
            "-created_at"
        )

        serializer = DICOMImageSerializer(
            page_queryset, many=True,
            context={'institute_name_map': _build_center_name_to_institute_map()}
        )
        return Response(
            {
                "success": True,
                "results": serializer.data,
                "institute_name": institute_name,
                "centers": center_names_list,
                "total_count": total_patients,
                "total_pages": total_pages,
                "current_page": page,
            }
        )
    except Exception as e:
        logger.error(f"Error in get_institute_studies: {str(e)}")
        return Response(
            {"success": False, "error": str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@api_view(["GET"])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def get_institute_stats(request):
    try:
        user = request.user
        profile = UserProfile.objects.filter(user=user).first()
        if not profile or not profile.center:
            return Response(
                {"success": False, "error": "No institute assigned"},
                status=status.HTTP_404_NOT_FOUND,
            )
        institute_name = profile.center.institute_name
        centers_in_institute = Center.objects.filter(institute_name=institute_name)
        center_names_list = []
        for center_obj in centers_in_institute:
            for center_name in center_obj.center_names.all():
                center_names_list.append(center_name.name)
        queryset = DICOMImage.objects.filter(center_name__in=center_names_list)
        total_images = queryset.count()
        total_studies = queryset.values("study_instance_uid").distinct().count()
        total_patients = queryset.values("patient_id").distinct().count()
        status_counts = queryset.values("status").annotate(count=Count("id"))
        total_size = queryset.aggregate(total=Sum("file_size"))["total"] or 0
        center_stats = []
        for center_name in center_names_list:
            center_queryset = queryset.filter(center_name=center_name)
            center_stats.append(
                {
                    "center_name": center_name,
                    "image_count": center_queryset.count(),
                    "patient_count": center_queryset.values("patient_id")
                    .distinct()
                    .count(),
                    "study_count": center_queryset.values("study_instance_uid")
                    .distinct()
                    .count(),
                }
            )
        return Response(
            {
                "success": True,
                "institute_name": institute_name,
                "total_centers": len(center_names_list),
                "total_images": total_images,
                "total_studies": total_studies,
                "total_patients": total_patients,
                "total_size_bytes": total_size,
                "total_size_mb": (
                    round(total_size / (1024 * 1024), 2) if total_size else 0
                ),
                "status_breakdown": {
                    item["status"]: item["count"] for item in status_counts
                },
                "center_stats": center_stats,
            }
        )
    except Exception as e:
        logger.error(f"Error in get_institute_stats: {str(e)}")
        return Response(
            {"success": False, "error": str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


def template_manager_page(request):
    from django.http import HttpResponse
    import os
    from django.conf import settings

    html_path = os.path.join(settings.BASE_DIR, "myapp", "static", "template.html")
    try:
        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()
        
        response = HttpResponse(html_content)
        
        if request.user.is_authenticated:
            token = Token.objects.get_or_create(user=request.user)[0]
            response.set_cookie('auth_token', token.key, httponly=False, samesite='Lax')
        
        return response
    except FileNotFoundError:
        return HttpResponse(
            f"<h1>Template Not Found</h1><p>File not found at: {html_path}</p>",
            status=404,
        )

@api_view(['GET'])
@permission_classes([AllowAny])
def get_report_templates(request):
    try:
        user = request.user
        if not user.is_authenticated:
            return Response(
                {"success": False, "error": "Authentication required"},
                status=status.HTTP_401_UNAUTHORIZED,
            )
        body_part = request.GET.get("body_part", "")
        modality = request.GET.get("modality", "")
        profile = UserProfile.objects.filter(user=user).first()
        is_admin = False
        if profile and profile.role:
            is_admin = profile.role.name == "Admin"
        elif user.is_superuser:
            is_admin = True
        
        if is_admin:
            templates = ReportTemplate.objects.filter(is_active=True)
        else:
            user_role = profile.role.name if profile and profile.role else None
            
            if user_role == 'Doctor':
                admin_group = Group.objects.filter(name="Admin").first()
                admin_profiles = UserProfile.objects.filter(role=admin_group) if admin_group else []
                admin_user_ids = [p.user_id for p in admin_profiles]
                superuser_ids = list(User.objects.filter(is_superuser=True).values_list("id", flat=True))
                
                subadmin_group = Group.objects.filter(name="SubAdmin").first()
                subadmin_profiles = UserProfile.objects.filter(role=subadmin_group) if subadmin_group else []
                subadmin_user_ids = [p.user_id for p in subadmin_profiles]
                
                center_group = Group.objects.filter(name="Center").first()
                center_profiles = UserProfile.objects.filter(role=center_group) if center_group else []
                center_user_ids = [p.user_id for p in center_profiles]
                
                privileged_user_ids = list(set(admin_user_ids + superuser_ids + subadmin_user_ids + center_user_ids))
                
                templates = ReportTemplate.objects.filter(is_active=True).filter(
                    Q(created_by=user) |
                    Q(created_by__in=privileged_user_ids) |
                    Q(created_by__isnull=True)
                )
            elif user_role in ['SubAdmin', 'Center']:
                templates = ReportTemplate.objects.filter(is_active=True).filter(
                    Q(created_by=user) |
                    Q(created_by__isnull=True)
                )
            else:
                admin_group = Group.objects.filter(name="Admin").first()
                admin_profiles = UserProfile.objects.filter(role=admin_group) if admin_group else []
                admin_user_ids = [p.user_id for p in admin_profiles]
                superuser_ids = list(User.objects.filter(is_superuser=True).values_list("id", flat=True))
                all_admin_ids = list(set(admin_user_ids + superuser_ids))
                templates = ReportTemplate.objects.filter(is_active=True).filter(
                    Q(created_by=user) |
                    Q(created_by__in=all_admin_ids) |
                    Q(created_by__isnull=True)
                )
        
        if body_part:
            templates = templates.filter(body_part__iexact=body_part)
        if modality:
            templates = templates.filter(modality__iexact=modality)
        templates = templates.order_by("body_part", "template_name")
        serializer = ReportTemplateSerializer(templates, many=True)
        grouped = {}
        for template in serializer.data:
            bp = template["body_part"]
            if bp not in grouped:
                grouped[bp] = []
            grouped[bp].append(template)
        return Response(
            {
                "success": True,
                "templates": serializer.data,
                "grouped": grouped,
                "count": templates.count(),
                "is_admin": is_admin,
            }
        )
    except Exception as e:
        logger.error(f"Error getting templates: {str(e)}")
        return Response(
            {"success": False, "error": str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


@api_view(['GET', 'POST', 'PUT', 'DELETE'])
@permission_classes([AllowAny])
def manage_templates(request):
    user = request.user
    if not user.is_authenticated:
        return Response(
            {"success": False, "error": "Authentication required"},
            status=status.HTTP_401_UNAUTHORIZED,
        )
    
    profile = UserProfile.objects.filter(user=user).first()
    is_admin = (
        profile and profile.role and profile.role.name == "Admin"
    ) or user.is_superuser
    
    can_manage = check_user_permission(user, 'manage_templates')
    
    if request.method == "GET":
        try:
            body_part = request.GET.get("body_part", "")
            modality = request.GET.get("modality", "")
            search = request.GET.get("search", "").strip()
            page = request.GET.get("page", "1")
            page_size = request.GET.get("page_size", "20")
            try:
                page = int(page)
            except ValueError:
                page = 1
            try:
                page_size = int(page_size)
            except ValueError:
                page_size = 20
            page_size = max(1, min(page_size, 200))
            
            if is_admin:
                templates = ReportTemplate.objects.filter(is_active=True)
            else:
                user_role = profile.role.name if profile and profile.role else None
                
                if user_role == 'Doctor':
                    admin_group = Group.objects.filter(name="Admin").first()
                    admin_profiles = UserProfile.objects.filter(role=admin_group) if admin_group else []
                    admin_user_ids = [p.user_id for p in admin_profiles]
                    superuser_ids = list(User.objects.filter(is_superuser=True).values_list("id", flat=True))
                    
                    subadmin_group = Group.objects.filter(name="SubAdmin").first()
                    subadmin_profiles = UserProfile.objects.filter(role=subadmin_group) if subadmin_group else []
                    subadmin_user_ids = [p.user_id for p in subadmin_profiles]
                    
                    center_group = Group.objects.filter(name="Center").first()
                    center_profiles = UserProfile.objects.filter(role=center_group) if center_group else []
                    center_user_ids = [p.user_id for p in center_profiles]
                    
                    privileged_user_ids = list(set(admin_user_ids + superuser_ids + subadmin_user_ids + center_user_ids))
                    
                    templates = ReportTemplate.objects.filter(is_active=True).filter(
                        Q(created_by=user) |
                        Q(created_by__in=privileged_user_ids) |
                        Q(created_by__isnull=True)
                    )
                elif user_role in ['SubAdmin', 'Center']:
                    templates = ReportTemplate.objects.filter(is_active=True).filter(
                        Q(created_by=user) |
                        Q(created_by__isnull=True)
                    )
                else:
                    admin_group = Group.objects.filter(name="Admin").first()
                    admin_profiles = UserProfile.objects.filter(role=admin_group) if admin_group else []
                    admin_user_ids = [p.user_id for p in admin_profiles]
                    superuser_ids = list(User.objects.filter(is_superuser=True).values_list("id", flat=True))
                    all_admin_ids = list(set(admin_user_ids + superuser_ids))
                    templates = ReportTemplate.objects.filter(is_active=True).filter(
                        Q(created_by=user) |
                        Q(created_by__in=all_admin_ids) |
                        Q(created_by__isnull=True)
                    )
            
            if body_part:
                templates = templates.filter(body_part__iexact=body_part)
            if modality:
                templates = templates.filter(modality__iexact=modality)
            if search:
                templates = templates.filter(
                    Q(template_name__icontains=search) | Q(body_part__icontains=search)
                )
            templates = templates.order_by("body_part", "template_name")
            
            total_count = templates.count()
            from django.core.paginator import Paginator
            paginator = Paginator(templates, page_size)
            page_obj = paginator.get_page(page)
            
            serializer = ReportTemplateSerializer(page_obj.object_list, many=True)
            return Response(
                {
                    "success": True,
                    "templates": serializer.data,
                    "count": total_count,
                    "page": page_obj.number,
                    "page_size": page_size,
                    "total_pages": paginator.num_pages,
                    "is_admin": is_admin,
                    "can_manage": can_manage
                }
            )
        except Exception as e:
            logger.error(f"Error retrieving templates: {str(e)}")
            return Response(
                {"success": False, "error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )
    elif request.method == "POST":
        if not can_manage:
            return Response(
                {"success": False, "error": "You don't have permission to create templates"},
                status=status.HTTP_403_FORBIDDEN
            )
        
        try:
            body_part = request.data.get("body_part", "").strip()
            modality = request.data.get("modality", "").strip().upper()
            template_name = request.data.get("template_name", "").strip()
            content = request.data.get("content", "").strip()
            is_active = request.data.get("is_active", True)
            if not body_part or not modality or not template_name or not content:
                return Response(
                    {
                        "success": False,
                        "error": "Body part, modality, name, and content required",
                    },
                    status=status.HTTP_400_BAD_REQUEST,
                )
            valid_body_parts = [c[0] for c in ReportTemplate.BODY_PART_CHOICES]
            valid_modalities = [c[0] for c in ReportTemplate.MODALITY_CHOICES]
            if body_part not in valid_body_parts:
                return Response(
                    {"success": False, "error": "Invalid body part"},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            if modality not in valid_modalities:
                return Response(
                    {"success": False, "error": "Invalid modality"},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            template = ReportTemplate.objects.create(
                body_part=body_part,
                modality=modality,
                template_name=template_name,
                content=content,
                is_active=is_active,
                created_by=user,
            )
            serializer = ReportTemplateSerializer(template)
            return Response(
                {
                    "success": True,
                    "message": "Template created",
                    "template": serializer.data,
                },
                status=status.HTTP_201_CREATED,
            )
        except Exception as e:
            logger.error(f"Error creating template: {str(e)}")
            return Response(
                {"success": False, "error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )
    elif request.method == "PUT":
        if not can_manage:
            return Response(
                {"success": False, "error": "You don't have permission to edit templates"},
                status=status.HTTP_403_FORBIDDEN
            )
        
        try:
            template_id = request.data.get("id")
            if not template_id:
                return Response(
                    {"success": False, "error": "Template ID required"},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            template = ReportTemplate.objects.get(id=template_id)
            if template.created_by and template.created_by != user and not is_admin:
                return Response(
                    {"success": False, "error": "No permission to edit"},
                    status=status.HTTP_403_FORBIDDEN,
                )
            new_body_part = request.data.get("body_part", template.body_part).strip()
            new_modality = request.data.get("modality", template.modality).strip().upper()
            valid_body_parts = [c[0] for c in ReportTemplate.BODY_PART_CHOICES]
            valid_modalities = [c[0] for c in ReportTemplate.MODALITY_CHOICES]
            if new_body_part not in valid_body_parts:
                return Response(
                    {"success": False, "error": "Invalid body part"},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            if not new_modality or new_modality not in valid_modalities:
                return Response(
                    {"success": False, "error": "Invalid modality"},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            template.body_part = new_body_part
            template.modality = new_modality
            template.template_name = request.data.get(
                "template_name", template.template_name
            ).strip()
            template.content = request.data.get("content", template.content).strip()
            template.is_active = request.data.get("is_active", template.is_active)
            template.save()
            serializer = ReportTemplateSerializer(template)
            return Response(
                {
                    "success": True,
                    "message": "Template updated",
                    "template": serializer.data,
                }
            )
        except ReportTemplate.DoesNotExist:
            return Response(
                {"success": False, "error": "Template not found"},
                status=status.HTTP_404_NOT_FOUND,
            )
        except Exception as e:
            logger.error(f"Error updating template: {str(e)}")
            return Response(
                {"success": False, "error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )
    elif request.method == "DELETE":
        if not can_manage:
            return Response(
                {"success": False, "error": "You don't have permission to delete templates"},
                status=status.HTTP_403_FORBIDDEN
            )
        
        try:
            template_id = request.data.get("id")
            if not template_id:
                return Response(
                    {"success": False, "error": "Template ID required"},
                    status=status.HTTP_400_BAD_REQUEST,
                )
            template = ReportTemplate.objects.get(id=template_id)
            if template.created_by and template.created_by != user and not is_admin:
                return Response(
                    {"success": False, "error": "No permission to delete"},
                    status=status.HTTP_403_FORBIDDEN,
                )
            template.is_active = False
            template.save()
            return Response({"success": True, "message": "Template deleted"})
        except ReportTemplate.DoesNotExist:
            return Response(
                {"success": False, "error": "Template not found"},
                status=status.HTTP_404_NOT_FOUND,
            )
        except Exception as e:
            logger.error(f"Error deleting template: {str(e)}")
            return Response(
                {"success": False, "error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_doctor_accessible_images(request):
    try:
        user = request.user
        profile = UserProfile.objects.filter(user=user).first()
        
        if not profile or not profile.role or profile.role.name != 'Doctor':
            return Response(
                {"success": False, "error": "Not a doctor account"},
                status=status.HTTP_403_FORBIDDEN
            )
        
        assigned_institutions = profile.assigned_institutions.all()
        
        if not assigned_institutions:
            return Response(
                {"success": True, "images": [], "message": "No institutions assigned"},
                status=status.HTTP_200_OK
            )
        
        center_names = []
        for institution in assigned_institutions:
            for center_name in institution.center_names.all():
                center_names.append(center_name.name)
        
        images = DICOMImage.objects.filter(
            center_name__in=center_names
        ).order_by('-created_at')
        
        status_filter = request.GET.get('status')
        if status_filter:
            images = images.filter(status=status_filter)
        
        serializer = DICOMImageSerializer(images, many=True)
        
        return Response({
            "success": True,
            "images": serializer.data,
            "count": images.count(),
            "assigned_institutions": [inst.institute_name for inst in assigned_institutions],
            "accessible_centers": center_names
        })
        
    except Exception as e:
        logger.error(f"Error in get_doctor_accessible_images: {str(e)}")
        return Response(
            {"success": False, "error": str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_doctors_by_institution(request):
    try:
        institution_id = request.GET.get('institution_id')
        center_name = request.GET.get('center_name')
        
        if institution_id:
            institution = Center.objects.get(id=institution_id)
            doctor_profiles = institution.assigned_doctors.filter(
                role__name='Doctor'
            ).select_related('user')
        elif center_name:
            center_name_obj = CenterName.objects.filter(name=center_name).first()
            if center_name_obj:
                institution = center_name_obj.center
                doctor_profiles = institution.assigned_doctors.filter(
                    role__name='Doctor'
                ).select_related('user')
            else:
                return Response(
                    {"success": False, "error": "Center not found"},
                    status=status.HTTP_404_NOT_FOUND
                )
        else:
            doctor_profiles = UserProfile.objects.filter(
                role__name='Doctor'
            ).select_related('user')
        
        doctors = [
            {
                'name': profile.full_name or profile.user.username,
                'username': profile.user.username,
                'designation': profile.designation,
                'qualification': profile.qualification,
                'bmdc_reg_no': profile.bmdc_reg_no,
                'assigned_institutions': [
                    inst.institute_name 
                    for inst in profile.assigned_institutions.all()
                ]
            }
            for profile in doctor_profiles
        ]
        
        return Response({'success': True, 'doctors': doctors})
        
    except Center.DoesNotExist:
        return Response(
            {"success": False, "error": "Institution not found"},
            status=status.HTTP_404_NOT_FOUND
        )
    except Exception as e:
        logger.error(f"Error in get_doctors_by_institution: {str(e)}")
        return Response(
            {"success": False, "error": str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        ) 

def check_user_permission(user, permission_type):
    if not user or not user.is_authenticated:
        return False
    
    if user.is_superuser:
        return True
    
    try:
        profile = UserProfile.objects.filter(user=user).first()
        if not profile:
            return False
        
        if permission_type == 'assign_doctors':
            return profile.can_assign_doctors
        elif permission_type == 'write_reports':
            return profile.can_write_reports
        elif permission_type == 'manage_templates':
            return profile.can_manage_templates
        
        return False
    except Exception as e:
        logger.error(f"Error checking permission: {str(e)}")
        return False    

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from django.http import HttpResponse
from rest_framework.decorators import api_view, permission_classes, authentication_classes
from rest_framework.permissions import AllowAny
from io import BytesIO
from bs4 import BeautifulSoup
import json
import os
from django.conf import settings

try:
    import qrcode
    QRCODE_AVAILABLE = True
except ImportError:
    QRCODE_AVAILABLE = False
    
@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def generate_report_docx(request):
    try:
        if not request.user or not request.user.is_authenticated:
            logger.error("DOCX Generation: User not authenticated")
            return HttpResponse(
                json.dumps({'success': False, 'error': 'Authentication required'}),
                status=401,
                content_type='application/json'
            )
        
        profile = UserProfile.objects.filter(user=request.user).first()
        if not profile or not profile.can_write_reports:
            if not request.user.is_superuser:
                logger.error(f"DOCX Generation: User {request.user.username} lacks permission")
                return HttpResponse(
                    json.dumps({'success': False, 'error': 'No permission to generate reports'}),
                    status=403,
                    content_type='application/json'
                )
        
        data = request.data
        logger.info(f"DOCX Generation started for user: {request.user.username}")
        
        patient_data = data.get('patient_data', {})
        report_html = data.get('report_content', '')
        doctor_info = data.get('doctor_info', {})
        patient_id = data.get('patient_id', 'Unknown')
        viewer_link = data.get('viewer_link', '')
        
        if not report_html or not report_html.strip():
            logger.error("DOCX Generation: No report content provided")
            return HttpResponse(
                json.dumps({'success': False, 'error': 'No report content provided'}),
                status=400,
                content_type='application/json'
            )
        
        logger.info(f"Generating DOCX for patient ID: {patient_id}")
        
        doc = Document()
        
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        header = section.header
        title_para = header.paragraphs[0]
        title_para.text = "DEPARTMENT OF RADIOLOGY & IMAGING"
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(13)
        title_run.font.name = 'Calibri'
        title_run.font.color.rgb = RGBColor(128, 128, 128)
        title_para_format = title_para.paragraph_format
        title_para_format.space_after = Pt(4)
        
        name = patient_data.get('patientName', 'Unknown')
        pid = patient_data.get('patientId', 'N/A')
        age = patient_data.get('patientAge', 'N/A')
        sex_value = patient_data.get('patientGender', 'U')
        sex = str(sex_value)[0].upper() if sex_value else 'U'
        doctor = patient_data.get('referringPhysician', 'N/A')
        exam_date = data.get('exam_date', 'N/A')
        template_name = data.get('template_name', '')
        report_submission_datetime = data.get('report_submission_datetime', 'N/A')
        
        tbl_xml = f'''<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:tblPr>
                <w:tblW w:w="5000" w:type="pct"/>
                <w:tblBorders>
                    <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                    <w:left w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                    <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                    <w:right w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                    <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                </w:tblBorders>
            </w:tblPr>
            <w:tr>
                <w:tc>
                    <w:tcPr><w:tcW w:w="1100" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:color w:val="000000"/></w:rPr>
                        <w:t>Name</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="150" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="center"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>:</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="2800" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>{name}</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="1300" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>Patient ID</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="150" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="center"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>:</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="2100" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>{pid}</w:t></w:r></w:p>
                </w:tc>
            </w:tr>
            <w:tr>
                <w:tc>
                    <w:tcPr><w:tcW w:w="1100" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>Age</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="150" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="center"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>:</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="2800" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>{age}Y</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="1300" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>Exam date</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="150" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="center"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>:</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="2100" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>{exam_date.split()[0] if exam_date != 'N/A' else 'N/A'}</w:t></w:r></w:p>
                </w:tc>
            </w:tr>
            <w:tr>
                <w:tc>
                    <w:tcPr><w:tcW w:w="1100" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>Sex</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="150" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="center"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>:</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="2800" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>{sex}</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="1300" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>Report date/Time</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="150" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="center"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>:</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="2100" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>{report_submission_datetime if report_submission_datetime != 'N/A' else 'N/A'}</w:t></w:r></w:p>
                </w:tc>
            </w:tr>
            <w:tr>
                <w:tc>
                    <w:tcPr><w:tcW w:w="1100" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>Refd By</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="150" w:type="dxa"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="center"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:b/><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>:</w:t></w:r></w:p>
                </w:tc>
                <w:tc>
                    <w:tcPr><w:tcW w:w="6350" w:type="dxa"/><w:gridSpan w:val="4"/><w:vAlign w:val="center"/></w:tcPr>
                    <w:p><w:pPr><w:jc w:val="left"/><w:spacing w:before="50" w:after="50"/></w:pPr>
                        <w:r><w:rPr><w:sz w:val="22"/><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/></w:rPr>
                        <w:t>{doctor}</w:t></w:r></w:p>
                </w:tc>
            </w:tr>
        </w:tbl>'''
        
        table_element = parse_xml(tbl_xml)
        header._element.append(table_element)
        
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(0)
        spacing_para.paragraph_format.space_after = Pt(12)
        
        if template_name and template_name.strip():
            template_heading = doc.add_paragraph()
            template_heading_run = template_heading.add_run(template_name)
            template_heading_run.bold = True
            template_heading_run.font.size = Pt(13)
            template_heading_run.font.name = 'Calibri'
            template_heading_run.underline = True
            template_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            template_heading.paragraph_format.space_after = Pt(8)
            template_heading.paragraph_format.space_before = Pt(4)
        
        soup = BeautifulSoup(report_html, 'html.parser')
        
        for element in soup.descendants:
            if element.name == 'p':
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.line_spacing = 1.2
                
                for child in element.children:
                    if child.name == 'strong' or child.name == 'b':
                        run = para.add_run(child.get_text())
                        run.bold = True
                        run.font.size = Pt(12)
                        run.font.name = 'Calibri'
                    elif child.name == 'em' or child.name == 'i':
                        run = para.add_run(child.get_text())
                        run.italic = True
                        run.font.size = Pt(12)
                        run.font.name = 'Calibri'
                    elif child.name == 'u':
                        run = para.add_run(child.get_text())
                        run.underline = True
                        run.font.size = Pt(12)
                        run.font.name = 'Calibri'
                    elif isinstance(child, str):
                        text_content = str(child)
                        if text_content.strip():
                            run = para.add_run(text_content)
                            run.font.size = Pt(12)
                            run.font.name = 'Calibri'
                        
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    para = doc.add_paragraph(li.get_text(), style='List Bullet')
                    para.runs[0].font.size = Pt(12)
                    para.runs[0].font.name = 'Calibri'
                    para.paragraph_format.space_after = Pt(3)
                    para.paragraph_format.space_before = Pt(0)
                    
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    para = doc.add_paragraph(li.get_text(), style='List Number')
                    para.runs[0].font.size = Pt(12)
                    para.runs[0].font.name = 'Calibri'
                    para.paragraph_format.space_after = Pt(3)
                    para.paragraph_format.space_before = Pt(0)
        
        doc.add_paragraph()

        def add_tight_para(document, text='', bold=False, font_size=11):
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            para = document.add_paragraph(style='No Spacing')
            pPr = para._p.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            spacing.set(qn('w:line'), '240')
            spacing.set(qn('w:lineRule'), 'auto')
            existing = pPr.find(qn('w:spacing'))
            if existing is not None:
                pPr.remove(existing)
            pPr.append(spacing)
            if text:
                run = para.add_run(text)
                run.bold = bold
                run.font.size = Pt(font_size)
                run.font.name = 'Calibri'
            return para

        add_tight_para(doc)

        signature_url = doctor_info.get('signature', '')
        if signature_url:
            try:
                sig_path = os.path.join(settings.MEDIA_ROOT, signature_url.lstrip('/media/'))
                if os.path.exists(sig_path):
                    doc.add_picture(sig_path, width=Inches(1.3))
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    pic_para = doc.paragraphs[-1]
                    pPr = pic_para._p.get_or_add_pPr()
                    spacing = OxmlElement('w:spacing')
                    spacing.set(qn('w:before'), '0')
                    spacing.set(qn('w:after'), '0')
                    existing = pPr.find(qn('w:spacing'))
                    if existing is not None:
                        pPr.remove(existing)
                    pPr.append(spacing)
                else:
                    logger.warning(f"Signature file not found: {sig_path}")
            except Exception as sig_error:
                logger.error(f"Error adding signature: {str(sig_error)}")

        add_tight_para(doc, doctor_info.get('full_name', ''), bold=True, font_size=12)

        if doctor_info.get('qualification'):
            for line in doctor_info.get('qualification').splitlines():
                line = line.strip()
                if line:
                    add_tight_para(doc, line, font_size=11)

        if doctor_info.get('designation'):
            for line in doctor_info.get('designation').splitlines():
                line = line.strip()
                if line:
                    add_tight_para(doc, line, font_size=11)

        if doctor_info.get('institute_name'):
            add_tight_para(doc, doctor_info.get('institute_name'), font_size=11)

        if doctor_info.get('bmdc_reg_no'):
            add_tight_para(doc, f"BMDC Reg No: {doctor_info.get('bmdc_reg_no')}", font_size=11)

        if doctor_info.get('mobile') or doctor_info.get('contact_number'):
            mobile = doctor_info.get('mobile') or doctor_info.get('contact_number')
            add_tight_para(doc, f"Mobile: {mobile}", font_size=11)
        
        if viewer_link and viewer_link.strip():
            try:
                if QRCODE_AVAILABLE:
                    import qrcode as qr_module
                    qr = qr_module.QRCode(
                        version=1,
                        error_correction=qr_module.constants.ERROR_CORRECT_L,
                        box_size=10,
                        border=2,
                    )
                    qr.add_data(viewer_link)
                    qr.make(fit=True)
                    
                    qr_img = qr.make_image(fill_color="black", back_color="white")
                    
                    qr_buffer = BytesIO()
                    qr_img.save(qr_buffer, format='PNG')
                    qr_buffer.seek(0)
                    
                    qr_para = doc.add_paragraph()
                    qr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    qr_run = qr_para.add_run()
                    qr_run.add_picture(qr_buffer, width=Inches(0.8))
                    qr_para.paragraph_format.space_before = Pt(0)
                    qr_para.paragraph_format.space_after = Pt(2)
                    
                    qr_text_para = doc.add_paragraph()
                    qr_text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    qr_text_run = qr_text_para.add_run("Scan to view images")
                    qr_text_run.font.size = Pt(9)
                    qr_text_run.font.name = 'Calibri'
                    qr_text_run.font.color.rgb = RGBColor(0, 102, 204)
                    qr_text_para.paragraph_format.space_before = Pt(0)
                    qr_text_para.paragraph_format.space_after = Pt(0)
                    
                    logger.info(f"QR code added successfully")
                else:
                    logger.warning("qrcode library not available")
            except Exception as qr_error:
                logger.error(f"Error generating QR code: {str(qr_error)}")
                traceback.print_exc()
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"DOCX generated successfully for patient ID: {patient_id}")
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename=Report_{patient_id}.docx'
        
        return response
        
    except Exception as e:
        logger.error(f"DOCX Generation Error: {str(e)}")
        logger.error(f"Error traceback: {traceback.format_exc()}")
        return HttpResponse(
            json.dumps({
                'success': False, 
                'error': f'Report generation failed: {str(e)}'
            }),
            status=500,
            content_type='application/json'
        )

from .models import ActiveViewer

@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def register_viewer(request):
    study_uid = request.data.get('study_uid')
    
    if not study_uid:
        return Response({
            'success': False,
            'error': 'study_uid is required'
        }, status=400)
    
    try:
        user_profile = UserProfile.objects.filter(user=request.user).first()
        user_name = user_profile.full_name if user_profile and user_profile.full_name else request.user.username
        
        ActiveViewer.cleanup_stale_viewers()
        
        viewer, created = ActiveViewer.objects.update_or_create(
            study_uid=study_uid,
            user=request.user,
            defaults={'user_name': user_name}
        )
        
        return Response({
            'success': True,
            'message': 'Viewer registered successfully'
        })
    except Exception as e:
        logger.error(f"Error registering viewer: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)

@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def update_viewer_activity(request):
    study_uid = request.data.get('study_uid')
    
    if not study_uid:
        return Response({
            'success': False,
            'error': 'study_uid is required'
        }, status=400)
    
    try:
        viewer = ActiveViewer.objects.filter(
            study_uid=study_uid,
            user=request.user
        ).first()
        
        if viewer:
            viewer.save()
            return Response({
                'success': True,
                'message': 'Activity updated'
            })
        else:
            return Response({
                'success': False,
                'error': 'Viewer not found'
            }, status=404)
    except Exception as e:
        logger.error(f"Error updating viewer activity: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)

@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def unregister_viewer(request):
    study_uid = request.data.get('study_uid')
    
    if not study_uid:
        return Response({
            'success': False,
            'error': 'study_uid is required'
        }, status=400)
    
    try:
        ActiveViewer.objects.filter(
            study_uid=study_uid,
            user=request.user
        ).delete()
        
        return Response({
            'success': True,
            'message': 'Viewer unregistered successfully'
        })
    except Exception as e:
        logger.error(f"Error unregistering viewer: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)

@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def check_active_viewers(request):
    study_uid = request.GET.get('study_uid')
    
    if not study_uid:
        return Response({
            'success': False,
            'error': 'study_uid is required'
        }, status=400)
    
    try:
        ActiveViewer.cleanup_stale_viewers()
        
        active_viewers = list(ActiveViewer.get_active_viewers(study_uid))
        
        current_user_profile = UserProfile.objects.filter(user=request.user).first()
        current_user_name = current_user_profile.full_name if current_user_profile and current_user_profile.full_name else request.user.username
        
        active_viewers = [viewer for viewer in active_viewers if viewer != current_user_name]
        
        return Response({
            'success': True,
            'active_viewers': active_viewers,
            'count': len(active_viewers)
        })
    except Exception as e:
        logger.error(f"Error checking active viewers: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)

@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def kick_all_viewers(request):
    study_uid = request.data.get('study_uid')
    
    if not study_uid:
        return Response({
            'success': False,
            'error': 'study_uid is required'
        }, status=400)
    
    try:
        current_user_profile = UserProfile.objects.filter(user=request.user).first()
        current_user_name = current_user_profile.full_name if current_user_profile and current_user_profile.full_name else request.user.username
        
        kicked_count = ActiveViewer.objects.filter(
            study_uid=study_uid
        ).exclude(user=request.user).delete()[0]
        
        return Response({
            'success': True,
            'message': f'Kicked {kicked_count} viewer(s)',
            'kicked_count': kicked_count
        })
    except Exception as e:
        logger.error(f"Error kicking viewers: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)

@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def am_i_active(request):
    study_uid = request.GET.get('study_uid')
    
    if not study_uid:
        return Response({
            'success': False,
            'error': 'study_uid is required'
        }, status=400)
    
    try:
        viewer = ActiveViewer.objects.filter(
            study_uid=study_uid,
            user=request.user
        ).first()
        
        return Response({
            'success': True,
            'is_active': viewer is not None
        })
    except Exception as e:
        logger.error(f"Error checking viewer status: {str(e)}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)

@api_view(["GET"])
@permission_classes([IsAuthenticated])
def get_image_report(request, dicom_id):
    try:
        dicom = DICOMImage.objects.get(id=dicom_id)
        report_file_url = None
        if dicom.report_file:
            report_file_url = request.build_absolute_uri(dicom.report_file.url)
        updated_at_str = ""
        if dicom.updated_at:
            dt = dicom.updated_at
            updated_at_str = dt.strftime('%d-%m-%Y %H:%M:%S')
        return Response({
            "success": True,
            "dicom_id": dicom.id,
            "report_content": dicom.report_content or "",
            "report_file": report_file_url,
            "status": dicom.status,
            "reported_by": dicom.reported_by,
            "updated_at": updated_at_str,
        })
    except DICOMImage.DoesNotExist:
        return Response({"success": False, "error": "DICOM not found"}, status=404)
    except Exception as e:
        return Response({"success": False, "error": str(e)}, status=500)


@api_view(["GET"])
@permission_classes([IsAuthenticated])
def get_study_reports(request):
    study_uid = request.query_params.get("study_uid", "")
    if not study_uid:
        return Response({"success": False, "error": "study_uid required"}, status=400)
    try:
        images = DICOMImage.objects.filter(study_instance_uid=study_uid).order_by("id")
        reports = []
        seen_files = set()

        for img in images:
            history = _read_history(img.id)

            if history:
                for entry in history:
                    rf = entry.get("report_file") or ""
                    if rf and rf not in seen_files:
                        seen_files.add(rf)
                        reports.append({
                            "dicom_id": img.id,
                            "report_content": entry.get("report_content", ""),
                            "report_file": rf,
                            "status": entry.get("status", ""),
                            "reported_by": entry.get("reported_by", ""),
                            "created_at": entry.get("created_at", ""),
                            "instance_number": img.instance_number,
                            "series_description": img.series_description,
                            "modality": img.modality,
                        })

                current_report_file_url = None
                if img.report_file:
                    try:
                        current_report_file_url = request.build_absolute_uri(img.report_file.url)
                    except:
                        pass
                if current_report_file_url and current_report_file_url not in seen_files:
                    if img.report_content and img.report_content.strip():
                        seen_files.add(current_report_file_url)
                        reports.append({
                            "dicom_id": img.id,
                            "report_content": img.report_content,
                            "report_file": current_report_file_url,
                            "status": img.status,
                            "reported_by": img.reported_by,
                            "created_at": "",
                            "instance_number": img.instance_number,
                            "series_description": img.series_description,
                            "modality": img.modality,
                        })
            else:
                if img.report_content and img.report_content.strip():
                    report_file_url = None
                    if img.report_file:
                        try:
                            report_file_url = request.build_absolute_uri(img.report_file.url)
                        except:
                            pass
                    if not report_file_url or report_file_url not in seen_files:
                        if report_file_url:
                            seen_files.add(report_file_url)
                        reports.append({
                            "dicom_id": img.id,
                            "report_content": img.report_content,
                            "report_file": report_file_url,
                            "status": img.status,
                            "reported_by": img.reported_by,
                            "created_at": "",
                            "instance_number": img.instance_number,
                            "series_description": img.series_description,
                            "modality": img.modality,
                        })

        return Response({"success": True, "reports": reports, "count": len(reports)})
    except Exception as e:
        return Response({"success": False, "error": str(e)}, status=500)




@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated, AdminPermission])
def update_user_profile(request):
    try:
        user_id = request.data.get('user_id')
        username = request.data.get('username')
        
        if user_id:
            try:
                user = User.objects.get(id=user_id)
            except User.DoesNotExist:
                return Response({'error': 'User not found'}, status=404)
        elif username:
            try:
                user = User.objects.get(username=username)
            except User.DoesNotExist:
                return Response({'error': 'User not found'}, status=404)
        else:
            return Response({'error': 'User ID or username required'}, status=400)
        
        profile, created = UserProfile.objects.get_or_create(user=user)
        
        if 'full_name' in request.data:
            profile.full_name = request.data.get('full_name')
        if 'designation' in request.data:
            profile.designation = request.data.get('designation')
        if 'qualification' in request.data:
            profile.qualification = request.data.get('qualification')
        if 'contact_number' in request.data:
            profile.contact_number = request.data.get('contact_number')
        if 'bmdc_reg_no' in request.data:
            profile.bmdc_reg_no = request.data.get('bmdc_reg_no')
        
        if 'signature' in request.FILES:
            profile.signature = request.FILES['signature']
        
        permission_overrides = {}
        if 'can_assign_doctors' in request.data:
            permission_overrides['can_assign_doctors'] = request.data.get('can_assign_doctors') in ['true', 'True', True, '1']
        if 'can_write_reports' in request.data:
            permission_overrides['can_write_reports'] = request.data.get('can_write_reports') in ['true', 'True', True, '1']
        if 'can_manage_templates' in request.data:
            permission_overrides['can_manage_templates'] = request.data.get('can_manage_templates') in ['true', 'True', True, '1']
        if 'can_view_images' in request.data:
            permission_overrides['can_view_images'] = request.data.get('can_view_images') in ['true', 'True', True, '1']
        
        if 'assigned_institutions' in request.data:
            try:
                assigned_inst_data = request.data.get('assigned_institutions')
                if isinstance(assigned_inst_data, str):
                    import json
                    assigned_inst_ids = json.loads(assigned_inst_data)
                else:
                    assigned_inst_ids = assigned_inst_data
                
                profile.assigned_institutions.clear()
                for inst_id in assigned_inst_ids:
                    try:
                        center = Center.objects.get(id=int(inst_id))
                        profile.assigned_institutions.add(center)
                    except (Center.DoesNotExist, ValueError):
                        continue
            except Exception as e:
                logger.error(f"Error updating assigned institutions: {e}")
        
        profile.save()
        
        if permission_overrides:
            UserProfile.objects.filter(pk=profile.pk).update(**permission_overrides)
        
        return Response({'success': True, 'message': 'Profile updated successfully'})
    
    except Exception as e:
        logger.error(f"Error updating user profile: {str(e)}")
        return Response({'error': str(e)}, status=500)


@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated, AdminPermission])
def get_user_assigned_institutions(request, user_id):
    try:
        user = User.objects.get(id=user_id)
        profile = UserProfile.objects.filter(user=user).first()
        
        if not profile:
            return Response({'assigned_institutions': []})
        
        assigned_inst_ids = list(profile.assigned_institutions.values_list('id', flat=True))
        
        return Response({'assigned_institutions': assigned_inst_ids})
    
    except User.DoesNotExist:
        return Response({'error': 'User not found'}, status=404)
    except Exception as e:
        logger.error(f"Error fetching assigned institutions: {str(e)}")
        return Response({'error': str(e)}, status=500)


@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated, AdminPermission])
def create_center_name(request):
    try:
        center_id = request.data.get('center')
        name = request.data.get('name')
        
        if not center_id or not name:
            return Response({'error': 'Center ID and name are required'}, status=400)
        
        try:
            center = Center.objects.get(id=center_id)
        except Center.DoesNotExist:
            return Response({'error': 'Center not found'}, status=404)
        
        center_name = CenterName.objects.create(
            center=center,
            name=name
        )
        
        return Response({
            'id': center_name.id,
            'name': center_name.name,
            'center': center_id
        }, status=201)
    
    except Exception as e:
        logger.error(f"Error creating center name: {str(e)}")
        return Response({'error': str(e)}, status=500)


@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def get_centers_with_doctors(request):
    try:
        centers = Center.objects.all()
        centers_data = []
        
        for center in centers:
            assigned_doctors = center.assigned_doctors.all()
            doctor_names = [
                doc.full_name or doc.user.username 
                for doc in assigned_doctors
            ]
            
            center_data = {
                'id': center.id,
                'institute_name': center.institute_name,
                'center_names': [cn.name for cn in center.center_names.all()],
                'is_default': center.is_default,
                'user': center.user.id if center.user else None,
                'assigned_doctors_names': doctor_names
            }
            centers_data.append(center_data)
        
        return Response(centers_data)
    
    except Exception as e:
        logger.error(f"Error fetching centers with doctors: {str(e)}")
        return Response({'error': str(e)}, status=500)


        

@api_view(['GET'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def list_active_sessions(request):
    try:
        tokens = Token.objects.select_related('user__userprofile__role', 'user__session_info').all()
        sessions = []
        for token in tokens:
            user = token.user
            session_info = getattr(user, 'session_info', None)
            if not session_info or not session_info.login_time:
                continue
            profile = getattr(user, 'userprofile', None)
            role = profile.role.name if profile and profile.role else 'No Role'
            login_time = session_info.login_time.isoformat()
            last_activity = session_info.last_activity.isoformat() if session_info.last_activity else None
            client_ip = session_info.client_ip
            sessions.append({
                'user_id': user.id,
                'username': user.username,
                'role': role,
                'login_time': login_time,
                'last_activity': last_activity,
                'client_ip': client_ip or '—',
            })
        return Response({'sessions': sessions})
    except Exception as e:
        return Response({'error': str(e)}, status=500)


@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def kill_session(request):
    try:
        user_id = request.data.get('user_id')
        if not user_id:
            return Response({'error': 'user_id is required'}, status=400)
        if int(user_id) == request.user.id:
            return Response({'error': 'Cannot kill your own session'}, status=400)
        Token.objects.filter(user_id=user_id).delete()
        UserSession.objects.filter(user_id=user_id).delete()
        return Response({'success': True})
    except Exception as e:
        return Response({'error': str(e)}, status=500)




import platform
import socket
import psutil
from datetime import datetime
from django.http import JsonResponse
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def server_info(request):
    uname = platform.uname()
    boot_ts = psutil.boot_time()
    boot_time = datetime.fromtimestamp(boot_ts).strftime('%Y-%m-%d %H:%M:%S')

    cpu_freq = psutil.cpu_freq()
    cpu_usage = psutil.cpu_percent(interval=0.5)

    mem = psutil.virtual_memory()

    def gb(b):
        return round(b / (1024 ** 3), 2)

    drives = []
    for part in psutil.disk_partitions(all=False):
        try:
            usage = psutil.disk_usage(part.mountpoint)
            total = gb(usage.total)
            used = gb(usage.used)
            free = gb(usage.free)
            used_pct = round(usage.used / usage.total * 100, 2) if usage.total else 0
            free_pct = round(100 - used_pct, 2)
            drives.append({
                'device': part.device,
                'mountpoint': part.mountpoint,
                'fstype': part.fstype,
                'total_gb': total,
                'used_gb': used,
                'free_gb': free,
                'used_percent': used_pct,
                'free_percent': free_pct,
            })
        except PermissionError:
            continue

    data = {
        'os_name': f"{uname.system} {uname.release}",
        'os_version': uname.version,
        'hostname': socket.gethostname(),
        'architecture': uname.machine,
        'boot_time': boot_time,

        'cpu_model': uname.processor or platform.processor(),
        'cpu_physical_cores': psutil.cpu_count(logical=False),
        'cpu_logical_cores': psutil.cpu_count(logical=True),
        'cpu_max_freq': round(cpu_freq.max, 1) if cpu_freq else None,
        'cpu_usage': cpu_usage,

        'memory_total_gb': gb(mem.total),
        'memory_used_gb': gb(mem.used),
        'memory_free_gb': gb(mem.free),
        'memory_available_gb': gb(mem.available),
        'memory_used_percent': mem.percent,

        'drives': drives,
    }

    return JsonResponse(data)


@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated, AdminPermission])
def change_user_password(request):
    try:
        user_id = request.data.get('user_id')
        new_password = request.data.get('new_password')
        if not user_id or not new_password:
            return Response({'error': 'user_id and new_password required'}, status=400)
        user = User.objects.get(id=user_id)
        user.set_password(new_password)
        user.save()
        return Response({'success': True})
    except User.DoesNotExist:
        return Response({'error': 'User not found'}, status=404)
    except Exception as e:
        return Response({'error': str(e)}, status=500)


@api_view(['POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated, AdminPermission])
def toggle_user_active(request):
    try:
        user_id = request.data.get('user_id')
        is_active = request.data.get('is_active')
        if user_id is None or is_active is None:
            return Response({'error': 'user_id and is_active required'}, status=400)
        user = User.objects.get(id=user_id)
        user.is_active = bool(is_active)
        user.save()
        return Response({'success': True, 'is_active': user.is_active})
    except User.DoesNotExist:
        return Response({'error': 'User not found'}, status=404)
    except Exception as e:
        return Response({'error': str(e)}, status=500)

import json as _json
import os as _os
import uuid as _uuid
from django.conf import settings as _settings
from django.utils import timezone as _timezone


def _get_groups_file():
    base = getattr(_settings, 'MEDIA_ROOT', '/tmp')
    return _os.path.join(base, 'doctor_groups.json')


def _read_doctor_groups():
    path = _get_groups_file()
    try:
        if _os.path.exists(path):
            with open(path, 'r') as f:
                return _json.load(f)
    except Exception:
        pass
    return []


def _write_doctor_groups(groups):
    path = _get_groups_file()
    try:
        with open(path, 'w') as f:
            _json.dump(groups, f, indent=2)
        return True
    except Exception:
        return False


@api_view(['GET', 'POST'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def doctor_groups_list(request):
    if request.method == 'GET':
        groups = _read_doctor_groups()
        return Response(groups)

    if request.method == 'POST':
        name = (request.data.get('name') or '').strip()
        description = (request.data.get('description') or '').strip()
        enabled = bool(request.data.get('enabled', True))
        if not name:
            return Response({'error': 'Name is required.'}, status=400)
        groups = _read_doctor_groups()
        new_group = {
            'id': 'grp_' + str(_uuid.uuid4()).replace('-', '')[:12],
            'name': name,
            'description': description,
            'enabled': enabled,
            'created_at': _timezone.now().isoformat(),
        }
        groups.append(new_group)
        _write_doctor_groups(groups)
        return Response(new_group, status=201)


@api_view(['GET', 'PUT', 'PATCH', 'DELETE'])
@authentication_classes([TokenAuthentication])
@permission_classes([IsAuthenticated])
def doctor_groups_detail(request, group_id):
    groups = _read_doctor_groups()
    idx = next((i for i, g in enumerate(groups) if g['id'] == group_id), None)
    if idx is None:
        return Response({'error': 'Group not found.'}, status=404)

    if request.method == 'GET':
        return Response(groups[idx])

    if request.method in ('PUT', 'PATCH'):
        g = groups[idx]
        if 'name' in request.data:
            name = (request.data['name'] or '').strip()
            if not name:
                return Response({'error': 'Name is required.'}, status=400)
            g['name'] = name
        if 'description' in request.data:
            g['description'] = (request.data.get('description') or '').strip()
        if 'enabled' in request.data:
            g['enabled'] = bool(request.data['enabled'])
        if 'members' in request.data:
            raw = request.data['members']
            if isinstance(raw, list):
                g['members'] = [str(m).strip() for m in raw if str(m).strip()]
            else:
                g['members'] = []
        groups[idx] = g
        _write_doctor_groups(groups)
        return Response(g)

    if request.method == 'DELETE':
        groups.pop(idx)
        _write_doctor_groups(groups)
        return Response({'success': True})
